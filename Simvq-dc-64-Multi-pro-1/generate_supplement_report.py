"""
生成 SimVQ 项目遗漏分析补充报告（第三版）
聚焦前两份报告中未覆盖的新发现
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "SimVQ_遗漏问题补充报告_v3.docx")
doc = Document()

# 样式
style = doc.styles['Normal']
font = style.font; font.name = '微软雅黑'; font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.25

for lv in range(1,5):
    hs = doc.styles[f'Heading {lv}']
    hs.font.name = '微软雅黑'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if lv==1: hs.font.size=Pt(18); hs.font.color.rgb=RGBColor(0xC0,0x39,0x2B)
    elif lv==2: hs.font.size=Pt(14); hs.font.color.rgb=RGBColor(0x2C,0x3E,0x50)
    elif lv==3: hs.font.size=Pt(12); hs.font.color.rgb=RGBColor(0x34,0x49,0x5E)

def add_code(doc, code, fs=8.5):
    for line in code.strip().split('\n'):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name='Consolas'; run.font.size=Pt(fs)
        run.font.color.rgb=RGBColor(0x2D,0x2D,0x2D)
        p.paragraph_format.space_before=Pt(0)
        p.paragraph_format.space_after=Pt(0)
        p.paragraph_format.left_indent=Cm(0.6)
        p.paragraph_format.line_spacing=1.0

def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h
        for p in c.paragraphs:
            for r in p.runs: r.bold=True; r.font.size=Pt(9)
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c=t.rows[ri+1].cells[ci]; c.text=str(val)
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(9)
    if col_widths:
        for i,w in enumerate(col_widths):
            for row in t.rows: row.cells[i].width=Cm(w)
    doc.add_paragraph()

def add_badge(doc, level):
    colors = {'CRITICAL': RGBColor(0xC0,0x39,0x2B), 'HIGH': RGBColor(0xE6,0x7E,0x22), 'MEDIUM': RGBColor(0xF3,0x9C,0x12), 'LOW': RGBColor(0x27,0xAE,0x60)}
    p=doc.add_paragraph()
    run=p.add_run(f'【{level}】')
    run.bold=True; run.font.color.rgb=colors.get(level,RGBColor(0,0,0)); run.font.size=Pt(11)

# ===== 封面 =====
doc.add_paragraph()
t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=t.add_run('SimVQ 遗漏问题补充报告'); r.bold=True; r.font.size=Pt(24); r.font.color.rgb=RGBColor(0xC0,0x39,0x2B)
doc.add_paragraph()
sub=doc.add_paragraph(); sub.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=sub.add_run('—— 前两份报告未覆盖的深层问题 ——'); r.font.size=Pt(13); r.font.color.rgb=RGBColor(0x7F,0x8C,0x8D)
doc.add_paragraph()
dt=doc.add_paragraph(); dt.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=dt.add_run('版本 V3.0  |  2026年5月14日  |  与前两版互补阅读'); r.font.size=Pt(10); r.font.color.rgb=RGBColor(0x95,0xA5,0xA6)
doc.add_page_break()

# ===== 概述 =====
doc.add_heading('关于本报告', level=1)
doc.add_paragraph(
    '本报告是前两份报告（SimVQ性能分析报告.docx 和 SimVQ_深度性能分析报告_v2.docx）的补充。'
    '前两份报告已经覆盖了 14 个核心问题（训练/测试信道不一致、STE梯度、BN失效、SimVQ局限性等）。'
    '本报告专门聚焦前两份报告中未涉及的新发现问题，按照对性能影响的严重程度排列。'
    '建议三份报告配合阅读。'
)
doc.add_paragraph(
    '经过对全部 17 个源文件的逐行重新审查，额外发现以下新问题。部分问题的影响程度不亚于之前报告中的 CRITICAL 级别问题。'
)

doc.add_page_break()

# ================================================================
doc.add_heading('新问题 1：验证管线同样使用错误信道模型 → 最优模型从未被正确选出', level=1)
add_badge(doc, 'CRITICAL')

doc.add_paragraph('【代码位置】train.py:186-202 (验证循环) + models/deepsc.py:90-122 (forward_val)')

doc.add_paragraph(
    '这是一个之前被完全忽略的系统性缺陷。训练脚本在每个 epoch 结束后运行验证，选择 val_loss 最低的 checkpoint 作为 best model。'
    '但验证过程使用的是 forward_val，后者同样调用 FiniteBlocklengthChannel！'
)
add_code(doc, '''# train.py:187-200 — 验证循环
deepsc_model.eval()
with torch.no_grad():
    for real_images in val_dataloader:
        out = deepsc_model.forward_val(real_images)  # ← 使用训练信道模型!
        recon_loss_val, _ = deepsc_loss_fn(real_images, out["reconstructed_images"], out["vq_losses"])
        val_loss_sum += recon_loss_val.item()

# --- 模型选择 ---
if avg_val_loss < best_val_loss:
    best_val_loss = avg_val_loss
    torch.save(deepsc_model.state_dict(), "best_vq_deepsc.pth")  # ← "最优"模型''')

doc.add_paragraph('这导致了以下后果链：')

doc.add_paragraph(
    '1. 验证信道 = FiniteBlocklengthChannel（理论 BER 独立翻转）—— 与 test_real.py 的真实 AWGN+LDPC 链路完全不同。'
)
doc.add_paragraph(
    '2. 因此，"val_loss 最低" 的模型是在 "错误信道模型"下表现最好的模型，而不是在 "真实信道" 下表现最好的模型。'
)
doc.add_paragraph(
    '3. 这解释了为什么即使 val_loss 持续降低，test_real.py 的 MS-SSIM 可能并不相应提升 —— 两个指标衡量的根本是不同的东西。'
)
doc.add_paragraph(
    '4. 更致命的是：不同 epoch 的 val_loss 受到随机 SNR 采样的噪声影响（forward_val 中 SNR 也是 random.uniform(0,15)），'
    '导致模型选择进一步随机化。'
)

doc.add_paragraph(
    '这个问题的严重性不亚于训练/测试信道不一致（问题 P1）。即使修复了训练信道，如果验证仍然使用错误的信道，'
    'best model 的选取依然是盲目的。'
)
doc.add_paragraph(
    '修正方向：在验证循环中，每 N 个 epoch（如每 5 epoch）对少量验证集图像运行一次真实的 test_real.py 链路评估，'
    '用真实 MS-SSIM 而非训练信道的 val_loss 来选择 best model。'
)

doc.add_page_break()

# ================================================================
doc.add_heading('新问题 2：解码器输出无界 + 测试指标裁剪掩盖了溢出错误', level=1)
add_badge(doc, 'HIGH')

doc.add_paragraph('【代码位置】models/semantic_decoder.py:114 (final层) + utils/metrics.py:164-165 (clip)')

doc.add_paragraph(
    '解码器的最终输出层是一个没有激活函数的 ConvTranspose2d：'
)
add_code(doc, '''# semantic_decoder.py:114
self.final = nn.ConvTranspose2d(in_ch, out_channels, kernel_size=3, stride=1, padding=1)
# ← 没有 nn.Tanh(), nn.Sigmoid(), 或任何输出约束!''')

doc.add_paragraph(
    '这意味着解码器的原始输出可以是任意实数值 —— 正值可以远大于 1，负值可以远小于 -1。'
    '而训练图像的归一化范围是 [-1, 1]（经过 normalize(0.5, 0.5)）。'
    '训练时 MSE loss 会"软约束"输出接近 [-1, 1]，但没有任何硬约束。'
)

doc.add_paragraph('在 test_real.py 中，评估前将图像从 [-1,1] 转换到 [0,1]：')
add_code(doc, '''# test_real.py:160-161
img1 = (real_image + 1) / 2       # [-1,1] → [0,1], 值域正确
img2 = (reconstructed_images + 1) / 2  # 可能超出[0,1]!

# metrics.py:164-165  — MS-SSIM计算中
img1 = np.clip(img1, 0, 1)   # 裁剪掉超出范围的值!
img2 = np.clip(img2, 0, 1)''')

doc.add_paragraph('如果 reconstructed_images 的某个像素值为 1.5（超出了 [-1,1] 范围），则：')
doc.add_paragraph('• img2 = (1.5+1)/2 = 1.25 → 被 clip 到 1.0')
doc.add_paragraph('• 这个像素的"真实重建误差"丢失了 —— 它本应是 1.25，却被当作 1.0 来计算 PSNR/MS-SSIM')
doc.add_paragraph('• 如果 reconstructed_images = -1.5，则 img2 = -0.25 → clip 到 0.0 → 同样丢失')
doc.add_paragraph(
    '这意味着当前报告的 MS-SSIM 和 PSNR 值是**系统性地偏乐观**的。解码器的 overshoot/undershoot '
    '错误被 clip 操作掩盖了。真实的图像质量可能低于报告值。'
)

doc.add_paragraph(
    '修正方向：在解码器末尾添加 `nn.Tanh()`，硬约束输出在 [-1, 1] 范围内。'
    '这样做有两个好处：(1) 消除 overshoot/undershoot，(2) 训练时的 MSE loss 更精确地反映重建质量。'
    '注意：训练数据经过 normalize(0.5, 0.5) 后严格在 [-1, 1] 范围内，所以 Tanh 不会造成信息损失。'
)

doc.add_page_break()

# ================================================================
doc.add_heading('新问题 3：Adam β₁=0.5 加剧训练不稳定性', level=1)
add_badge(doc, 'HIGH')

doc.add_paragraph('【代码位置】config.py:26')

add_code(doc, '''BETAS = (0.5, 0.999)  # Adam的β₁极低! 默认值是(0.9, 0.999)''')

doc.add_paragraph(
    '这是一个容易被忽略但影响深远的超参数选择。Adam 的 β₁ 控制一阶矩（动量）的衰减速度：'
)
doc.add_paragraph('• 默认 β₁=0.9 → 半衰期 ≈ ln(0.5)/ln(0.9) ≈ 6.6 步 → 梯度平滑')
doc.add_paragraph('• 当前 β₁=0.5 → 半衰期 ≈ ln(0.5)/ln(0.5) = 1.0 步 → 几乎无平滑！')

doc.add_paragraph(
    'β₁=0.5 意味着优化器几乎不对梯度做动量平滑，每次更新几乎完全由当前 batch 的梯度决定。'
    '在这个项目中，每个 batch 的梯度已经天然包含高噪声：'
)
doc.add_paragraph('• SNR 从 [0, 15] dB 随机采样 → 不同 SNR 的梯度量级差异巨大')
doc.add_paragraph('• 信道 BER 随 SNR 剧烈变化 → 低 SNR 时梯度以"纠错"为主，高 SNR 时以"精修"为主')
doc.add_paragraph('• 码本投影层和编解码器参数的量级差异 → 梯度本来就跨数量级')

doc.add_paragraph(
    'β₁=0.5 使优化器完全暴露在这些噪声之下，每次都朝当前噪声梯度方向大幅更新。'
    '这导致：(a) 训练 loss 曲线剧烈抖动，(b) 参数在最优解附近震荡而无法精确收敛，'
    '(c) 对学习率极为敏感 —— 稍大就不稳定，稍小就不动。'
)

doc.add_paragraph(
    '推测：代码作者可能发现默认 β₁=0.9 时训练"太慢"，所以降低 β₁ 来加速响应。'
    '但这只是用不稳定性换取了表面上的收敛速度。建议改回 β₁=0.9，或者折中使用 β₁=0.8。'
    '如果担心收敛慢，可以通过提高学习率或使用学习率 warmup 来补偿。'
)

doc.add_page_break()

# ================================================================
doc.add_heading('新问题 4：编码器特征与码本向量的尺度不匹配 → 距离计算失真', level=1)
add_badge(doc, 'HIGH')

doc.add_paragraph('【代码位置】models/vector_quantizer.py:47-50')

add_code(doc, '''# 当前：原始欧氏距离
d = torch.sum(flat ** 2, dim=1, keepdim=True) + \
    torch.sum(embed_weight ** 2, dim=1) - 2 * \
    torch.einsum('bd,nd->bn', flat, embed_weight)

encoding_idx = torch.argmin(d, dim=1)  # 选距离最近的码字''')

doc.add_paragraph(
    '这个距离计算存在一个隐含假设：编码器特征和码本向量的 L2 范数具有可比性。'
    '但在这个项目中，这个假设经常不成立：'
)

doc.add_paragraph(
    '1. 编码器特征经过 PReLU（无上界）→ 范数可以很大'
)
doc.add_paragraph(
    '2. 码本向量 = P × e₀_k，其中 e₀_k 初始化范数 ≈ 1.0，P 初始化为 Kaiming uniform → 初始范数中等'
)
doc.add_paragraph(
    '3. 训练过程中，P 被优化以最小化量化误差 → 码本向量范数可能显著变化'
)
doc.add_paragraph(
    '4. 如果 ||码本向量|| ≫ ||编码器特征||，距离由 ||e||² 项主导 → 所有特征被分配到范数最小的码字 → 码本坍缩到单个码字'
)
doc.add_paragraph(
    '5. 如果 ||编码器特征|| ≫ ||码本向量||，距离由 ||x||² 项主导 → 码字间的距离差异被淹没 → 分配变得随机'
)

doc.add_paragraph(
    '标准 VQ-VAE 的做法是在计算距离前将编码器特征和码本向量都 L2-归一化到单位球面上，'
    '从而距离等价于余弦相似度：'
)
add_code(doc, '''# 推荐：L2归一化后计算余弦距离
flat_norm = F.normalize(flat, dim=1)       # 编码器特征 → 单位向量
embed_norm = F.normalize(embed_weight, dim=1)  # 码本 → 单位向量
d = 2 - 2 * torch.mm(flat_norm, embed_norm.t())  # 余弦距离''')

doc.add_paragraph(
    'L2 归一化的额外好处：(a) 距离计算不受特征和码本的相对尺度影响，'
    '(b) 码本坍缩风险降低（因为码字间的角度关系比距离关系更稳定），'
    '(c) 训练更稳定（距离范围固定在 [0, 2]）。'
)

doc.add_page_break()

# ================================================================
doc.add_heading('新问题 5：commitment_cost=0.25 导致编码器-码本不对称优化', level=1)
add_badge(doc, 'MEDIUM')

doc.add_paragraph('【代码位置】models/vector_quantizer.py:60-62')

add_code(doc, '''e_latent_loss = F.mse_loss(quantized_bhwc.detach(), inputs_bhwc)  # 编码器承诺损失
q_latent_loss = F.mse_loss(quantized_bhwc, inputs_bhwc.detach())  # 码本学习损失
vq_loss = q_latent_loss + self.commitment_cost * e_latent_loss
#                          ↑ 0.25: 编码器承诺损失权重仅为码本损失的1/4''')

doc.add_paragraph(
    '传统 VQ-VAE 中 commitment_cost 通常设为 0.25 是有道理的 —— 编码器需要自由地学习好的表征，'
    '码本应该主动"追逐"编码器的输出。但在 SimVQ 中，情况完全不同：'
)

doc.add_paragraph(
    '• q_latent_loss 的梯度流向投影层 P（码本全局参数）。P 接收到的信号是"如何旋转/缩放码字使其接近编码器特征"。'
)
doc.add_paragraph(
    '• e_latent_loss 的梯度流向编码器。编码器接收到的信号是"如何改变特征使其接近码本"。由于 commitment_cost=0.25，这个信号被衰减了 4 倍。'
)
doc.add_paragraph(
    '• 在 SimVQ 中，码本的适应能力本就受到投影层的全局约束（只能做线性变换），如果同时编码器又不被充分激励去靠近码本，'
    '结果就是编码器特征逐渐偏离码本覆盖范围 → 量化误差持续增大。'
)

doc.add_paragraph(
    '建议：将 commitment_cost 增大到 0.5-1.0，或者采用动态调度（训练初期低值让编码器自由探索，'
    '后期高值让编码器严格承诺到码本）。'
)

doc.add_page_break()

# ================================================================
doc.add_heading('新问题 6：比特填充零值导致索引恢复偏向索引 0', level=1)
add_badge(doc, 'MEDIUM')

doc.add_paragraph('【代码位置】utils/bit_utils.py:33 + communications/ldpc_coding.py:49/79')

doc.add_paragraph(
    '在 LDPC 编码和译码过程中，都需要对输入比特进行块对齐填充（padding）：'
)
add_code(doc, '''# ldpc_coding.py:49 — 编码前填充
padded_bits = np.pad(bits, (0, padded_len - len(bits)), 'constant', constant_values=0)

# ldpc_coding.py:79 — 译码前填充 (LLR)
padded_llr = np.pad(received_llr, (0, padded_len - len(received_llr)), 'constant', constant_values=0.0)''')

doc.add_paragraph(
    '填充值 0 在 BPSK 映射中对应符号 -1（bit 0 → -1）。填充的比特经过 LDPC 编码后成为校验比特的一部分。'
    '在译码端，填充的 LLR 为 0.0 表示"完全不确定"（等概率为 0 或 1）。这在理想情况下是正确的。'
)

doc.add_paragraph('但在索引恢复阶段（bits_to_indices）：')
add_code(doc, '''# bit_utils.py:32-34
if len(scale_bits) < num_bits_for_scale:
    scale_bits = np.pad(scale_bits, (0, num_bits_for_scale - len(scale_bits)), 'constant')
    # ← 用 0 填充缺失的比特位，导致对应索引位偏向 0''')

doc.add_paragraph(
    '当 LDPC 译码后的比特数少于预期时（极少数情况），用 0 填充会将缺失索引的二进制位全部置 0，'
    '导致恢复出的索引系统地偏向 0（码本中第一个码字）。在高 SNR 下译码通常正确，这个问题不常触发。'
    '但在低 SNR 或译码失败的边界情况下，这会引入额外的索引偏差。'
)

doc.add_paragraph(
    '修正方向：当比特数不足时，填充随机比特而非零值，或者用 LLR=0 对应位置填充随机比特来消除偏差。'
    '更根本的方案是记录实际比特数并在编码前约定填充策略。'
)

doc.add_page_break()

# ================================================================
doc.add_heading('新问题 7：全局梯度裁剪对所有参数使用同一阈值', level=1)
add_badge(doc, 'MEDIUM')

doc.add_paragraph('【代码位置】train.py:164')

add_code(doc, '''torch.nn.utils.clip_grad_norm_(deepsc_model.parameters(), max_norm=1.0)''')

doc.add_paragraph(
    '这行代码对整个模型的所有参数（编码器 + 解码器 + 码本投影层）使用同一个 max_norm=1.0 进行梯度裁剪。'
    '但实际上，这三组参数的梯度天然具有不同的量级和噪声特性：'
)

doc.add_paragraph('• 编码器参数（~2M params, 12 层卷积）：梯度来自 MSE + 0.25×commitment loss → 量级中等')
doc.add_paragraph('• 解码器参数（~3M params, 8 层卷积 + skip）：梯度主要来自 MSE → 量级较大')
doc.add_paragraph('• 码本投影层参数（~3M params, 4 个 D×D 矩阵）：梯度从码本损失来 → 量级取决于码本大小和特征维度')

doc.add_paragraph(
    '由于解码器接收的是有噪特征（forward_train 中 quantized_noisy），'
    '解码器梯度可能比编码器梯度大一个数量级（因为需要纠正噪声引入的误差）。'
    '全局梯度裁剪意味着：解码器梯度被裁剪到 1.0，但同时编码器梯度可能只有 0.1，完全没被裁剪。'
    '这进一步加剧了编解码器学习速度的不平衡。'
)

doc.add_paragraph(
    '建议：使用分组梯度裁剪，为编码器、解码器和投影层分别设置不同的 max_norm。'
    '或者使用 `clip_grad_norm_` 的 `parameters` 参数分组传入。'
)

doc.add_page_break()

# ================================================================
doc.add_heading('新问题 8：SimVQ 投影层无正则化 → 有效秩可能退化', level=1)
add_badge(doc, 'MEDIUM')

doc.add_paragraph('【代码位置】models/vector_quantizer.py:17 + config.py')

doc.add_paragraph(
    'SimVQ 的投影层是一个 D×D 的无偏置线性层（Layer 3 为 1024×1024 = 1,048,576 个参数）。'
    '如深度分析报告（v2）第三部分所述，有效自由度仅约 K×D = 65,536 个。'
    '这意味着约 98% 的参数对码本几何没有贡献，但它们在训练中仍接收梯度并产生更新。'
)

doc.add_paragraph(
    '没有正则化（weight decay = 0）的情况下，这些"无效维度"的参数会随着训练不断漂移，可能导致：'
)
doc.add_paragraph(
    '1. 投影矩阵 P 的奇异值谱退化：有效秩从初始的 min(K,D)=64 逐渐下降，码本被压缩到更低的维度'
)
doc.add_paragraph(
    '2. 梯度噪声在冗余参数中累积，通过反向传播影响有效参数的学习'
)
doc.add_paragraph(
    '3. 当 P 的某些奇异值被训练得极小或极大时，码本向量在这些方向上的分量会坍缩或爆炸'
)

doc.add_paragraph(
    '建议：对投影层参数设置适度的 weight decay（如 1e-5 到 1e-4），'
    '定期监控 P 的奇异值谱（已在 v2 报告 3.4 节提供代码），'
    '或使用正交正则化鼓励 P 保持接近正交矩阵（有助于保持码字间的最小距离）。'
)

doc.add_page_break()

# ================================================================
doc.add_heading('新问题 9：训练/验证/测试的数据预处理不一致', level=1)
add_badge(doc, 'MEDIUM')

doc.add_paragraph('【代码位置】data/datasets.py:42-62')

doc.add_paragraph('三种模式的预处理策略不同：')
add_code(doc, '''# train: Resize(256) → RandomCrop(256)
# val:   Resize(256) → CenterCrop(256)
# test:  Resize((768, 512)) → 无Crop''')

doc.add_paragraph(
    '验证集（val）使用 CenterCrop(256)，训练集使用 RandomCrop(256)。两者都是从 Cars196 数据集加载的，'
    '但裁剪策略不同意味着验证集图像总是取中心区域，而训练集包含边缘区域。'
    '这个差异通常不重要（标准做法），但结合本项目中的数据域偏移问题（Cars196→Kodak），'
    '验证集不能完全代表训练分布，更不能代表测试分布。'
)

doc.add_paragraph(
    '测试集（test）使用 Kodak 数据集 + Resize((768, 512))。这里有一个细节：'
    'Kodak 原始图像是 768×512 或 512×768（竖向）。Resize((768, 512)) 会强制将竖向图像拉伸为横向，'
    '导致严重的宽高比失真。例如，一张 512×768 的竖向人像被 Resize 成 768×512 后，人物会被横向压扁。'
    '这种失真后的"原图"本身就是错误的参考图像，在此基础上计算 MS-SSIM/PSNR 没有意义。'
)

doc.add_paragraph(
    '建议：测试预处理应该保持原始宽高比 —— 使用 Resize(256) 或保持原尺寸，'
    '确保"原图"和"重建图"的宽高比一致。如果需要测试不同分辨率，'
    '应该按原始比例缩放而非强制变形。'
)

doc.add_page_break()

# ================================================================
doc.add_heading('新问题 10：学习率调度器从 checkpoint 恢复后可能已衰减至极低值', level=1)
add_badge(doc, 'LOW')

doc.add_paragraph('【代码位置】config.py:23-24 + train.py:89/99')

add_code(doc, '''LEARNING_RATE_G = 1.75e-5
# StepLR(step_size=100, gamma=0.5) → 每100个epoch学习率减半
# 恢复checkpoint时，scheduler状态也被恢复

# 如果从 epoch 250 恢复:
# LR = 1.75e-5 × 0.5^2 = 4.375e-6  (已经非常低)
# 如果从 epoch 350 恢复:
# LR = 1.75e-5 × 0.5^3 = 2.188e-6  (基本不更新)''')

doc.add_paragraph(
    '虽然当前 show 的训练已接近 400 epoch 的设定，但如果从较早的 checkpoint 恢复后再训练，'
    '学习率可能已经衰减到极低水平，导致参数几乎不更新。此时即使架构改动是正确的，'
    '也会因为 LR 太低而看不出效果，造成"改动无效"的误判。'
)
doc.add_paragraph(
    '建议：在恢复 checkpoint 后打印当前实际学习率。如果 value < 1e-6，考虑手动提高或重置 scheduler。'
)

doc.add_page_break()

# ================================================================
doc.add_heading('新问题 11：MS-SSIM 失败时返回 0.0 拉低平均值', level=1)
add_badge(doc, 'LOW')

doc.add_paragraph('【代码位置】utils/metrics.py:135-177 (calculate_ms_ssim)')

doc.add_paragraph(
    'calculate_ms_ssim 函数在遇到 NaN / Inf / 形状不匹配 / 异常时统一返回 0.0。'
    '虽然设计了多重防护，但如果某些特定图像（如极暗/极亮/单一色调）始终触发这些边界情况，'
    '它们的 0.0 值会被计入平均 MS-SSIM，拉低整体分数。这些失败可能不代表重建质量真的为 0，'
    '而是评估函数对边界情况的处理问题。'
)
doc.add_paragraph(
    '建议：记录哪些图像触发了 0.0 返回，单独分析它们的重建质量。'
    '对于合法的边界情况（如纯色图像），MS-SSIM 应该接近 1.0 而非 0.0。'
)

doc.add_page_break()

# ================================================================
doc.add_heading('新问题汇总', level=1)

add_table(doc,
    ['编号', '问题', '严重度', '代码位置', '对 test_real 的直接影响'],
    [
        ['N1', '验证使用错误信道 → 最优模型未被选出', 'CRITICAL', 'train.py:186-202\ndeepsc.py:90-122', '当前 best model 可能不是\n真正的最优模型'],
        ['N2', '解码器输出无界 + clip掩盖错误', 'HIGH', 'semantic_decoder.py:114\nmetrics.py:164-165', '报告的MS-SSIM/PSNR\n系统性偏高'],
        ['N3', 'Adam β₁=0.5 训练不稳定', 'HIGH', 'config.py:26', '参数震荡无法精确收敛'],
        ['N4', '编码器/码本尺度不匹配', 'HIGH', 'vector_quantizer.py:47-50', 'VQ量化质量受限'],
        ['N5', 'commitment_cost不对称', 'MEDIUM', 'vector_quantizer.py:60-62', '编码器偏离码本\n量化误差增大'],
        ['N6', '比特填充偏向索引0', 'MEDIUM', 'bit_utils.py:33\nldpc_coding.py:49/79', '低SNR下索引恢复偏差'],
        ['N7', '全局梯度裁剪阈值', 'MEDIUM', 'train.py:164', '编解码器学习不平衡'],
        ['N8', '投影层无正则化', 'MEDIUM', 'vector_quantizer.py:17', '投影矩阵秩退化'],
        ['N9', '预处理宽高比失真', 'MEDIUM', 'datasets.py:59', '测试参考图本身失真'],
        ['N10', '恢复后LR过低', 'LOW', 'train.py:89/99', '恢复训练后可能不更新'],
        ['N11', 'MS-SSIM失败返回0.0', 'LOW', 'metrics.py:135-177', '少数图像拉低平均值'],
    ],
    col_widths=[1, 3.5, 1.5, 3, 3.5]
)

doc.add_paragraph()
doc.add_paragraph(
    '以上 11 个新问题，加上 v1 报告中的 8 个问题和 v2 报告中的 14 个问题，'
    '共计识别出 33 个潜在性能影响因素。其中 CRITICAL/HIGH 级别的核心问题约 10 个，'
    '涵盖了信道模型、梯度流、归一化层、码本设计、优化器配置、损失函数、数据预处理等所有关键环节。'
)

doc.add_page_break()

# ================================================================
doc.add_heading('三版报告问题总索引', level=1)
doc.add_paragraph('为方便查阅，将三份报告的全部问题汇总如下：')

add_table(doc,
    ['ID', '问题', '严重度', '报告'],
    [
        ['P1', '训练/测试信道模型完全不同', 'CRITICAL', 'v1'],
        ['P2', 'STE梯度截断 → 编码器不学习鲁棒性', 'CRITICAL', 'v1'],
        ['P3', '调制方式训练/测试不匹配', 'CRITICAL', 'v1'],
        ['N1', '验证使用错误信道 → 最优模型未选出', 'CRITICAL', 'v3 ★新增'],
        ['P4', '仅用MSE损失 → 重建模糊', 'HIGH', 'v1'],
        ['P5', '码本容量不足 (K=64)', 'HIGH', 'v1'],
        ['P6', '训练/测试数据域偏移', 'HIGH', 'v1'],
        ['BN', '20层BN的6个失效模式 (v2详析)', 'HIGH', 'v2'],
        ['SV', 'SimVQ表达能力下界分析 (v2详析)', 'HIGH', 'v2'],
        ['N2', '解码器无界输出 + clip掩盖错误', 'HIGH', 'v3 ★新增'],
        ['N3', 'Adam β₁=0.5 训练不稳定', 'HIGH', 'v3 ★新增'],
        ['N4', '编码器/码本尺度不匹配', 'HIGH', 'v3 ★新增'],
        ['P7', 'VQ损失权重失衡', 'MEDIUM', 'v1'],
        ['P8', 'Skip Dropout gap', 'MEDIUM', 'v1'],
        ['P9', 'Nearest-neighbor块效应', 'MEDIUM', 'v2'],
        ['P10', 'LDPC译码失败无检测', 'MEDIUM', 'v2'],
        ['P11', 'FiniteBlocklengthChannel 5缺陷', 'MEDIUM', 'v2'],
        ['P12', 'Clamping码字偏差', 'MEDIUM', 'v2'],
        ['N5', 'commitment_cost不对称', 'MEDIUM', 'v3 ★新增'],
        ['N6', '比特填充偏向索引0', 'MEDIUM', 'v3 ★新增'],
        ['N7', '全局梯度裁剪', 'MEDIUM', 'v3 ★新增'],
        ['N8', '投影层无正则化', 'MEDIUM', 'v3 ★新增'],
        ['N9', '预处理宽高比失真', 'MEDIUM', 'v3 ★新增'],
        ['N10', '恢复后LR过低', 'LOW', 'v3 ★新增'],
        ['N11', 'MS-SSIM失败返回0.0', 'LOW', 'v3 ★新增'],
    ],
    col_widths=[1, 5, 1.5, 2]
)

# 保存
doc.save(OUTPUT_PATH)
print(f"补充报告已生成: {OUTPUT_PATH}")
print(f"文件大小: {os.path.getsize(OUTPUT_PATH)/1024:.1f} KB")
