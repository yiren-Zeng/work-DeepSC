"""
生成 SimVQ 项目性能分析报告 Word 文档
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "SimVQ性能分析报告.docx")

doc = Document()

# ===== 样式设置 =====
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.name = '微软雅黑'
    heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if level == 1:
        heading_style.font.size = Pt(18)
        heading_style.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    elif level == 2:
        heading_style.font.size = Pt(14)
        heading_style.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    elif level == 3:
        heading_style.font.size = Pt(12)
        heading_style.font.color.rgb = RGBColor(0x34, 0x49, 0x5E)

def add_code_block(doc, code_text):
    """添加代码块"""
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph(line)
        p.style = doc.styles['Normal']
        for run in p.runs:
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(1)

def add_table_simple(doc, headers, rows):
    """添加简单表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            table.rows[r + 1].cells[c].text = str(val)
    doc.add_paragraph()

def add_severity_badge(doc, level):
    """添加严重程度标记"""
    colors = {
        'CRITICAL': RGBColor(0xE7, 0x4C, 0x3C),
        'HIGH': RGBColor(0xE6, 0x7E, 0x22),
        'MEDIUM': RGBColor(0xF3, 0x9C, 0x12),
    }
    p = doc.add_paragraph()
    run = p.add_run(f"【严重程度: {level}】")
    run.bold = True
    run.font.color.rgb = colors.get(level, RGBColor(0, 0, 0))
    run.font.size = Pt(11)

# ============================
# 封面
# ============================
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('SimVQ 语义通信系统\n性能诊断与优化方案报告')
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

doc.add_paragraph()
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Simvq-dc-64-Multi-pro 项目')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run('2026年5月13日')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x95, 0xA5, 0xA6)

doc.add_page_break()

# ============================
# 目录提示
# ============================
doc.add_heading('报告目录', level=1)
toc_items = [
    '一、项目概述',
    '二、代码架构总览',
    '三、性能问题根因诊断（8个问题）',
    '   3.1 CRITICAL 级别',
    '   3.2 HIGH 级别',
    '   3.3 MEDIUM 级别',
    '四、优化方案（3套方案对比）',
    '五、推荐方案详细说明',
    '六、总结与建议',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ============================
# 一、项目概述
# ============================
doc.add_heading('一、项目概述', level=1)

doc.add_paragraph(
    '本项目是一个基于语义通信（Semantic Communication）与矢量量化（Vector Quantization）的'
    '端到端图像传输系统，名为 SimVQ（Semantic Image Vector Quantization）。系统目标是在'
    '带宽受限的无线信道（AWGN）条件下，实现高质量的图像重建。'
)

doc.add_heading('1.1 系统工作流程', level=2)
doc.add_paragraph(
    '发送端：原始图像 → 语义编码器（4层下采样）→ 各层 VQ 量化 → 索引转比特流 → LDPC 信道编码 → BPSK 调制 → 发送'
)
doc.add_paragraph(
    '接收端：接收符号 → LLR 软解调 → LDPC 译码 → 比特还原索引 → 语义解码器（4层上采样+跳跃连接）→ 重建图像'
)

doc.add_heading('1.2 核心配置参数', level=2)
add_table_simple(doc,
    ['参数', '值', '说明'],
    [
        ['输入尺寸', '256×256×3 (训练) / 768×512 (测试)', 'RGB图像'],
        ['下采样层数', '4', '每层stride=2'],
        ['码本大小 K', '[64, 64, 64, 64]', '各层码字数量'],
        ['嵌入维度 D', '[128, 256, 512, 1024]', '各层特征维度'],
        ['VQ损失权重', '[1, 1, 5, 10]', '深层权重更高'],
        ['学习率', '1.75e-5 / 1.75e-4(码本投影层)', 'Adam优化器'],
        ['训练SNR范围', '[0, 15] dB', '均匀随机采样'],
        ['训练信道模型', 'FiniteBlocklengthChannel', '理论BER近似'],
        ['测试信道', 'BPSK + AWGN + LDPC', '真实物理层链路'],
        ['损失函数', 'MSE + 加权VQ损失', '纯MSE重建'],
        ['训练数据集', 'Cars196', '汽车图像'],
        ['测试数据集', 'Kodak', '通用自然图像'],
    ]
)

doc.add_page_break()

# ============================
# 二、代码架构总览
# ============================
doc.add_heading('二、代码架构总览', level=1)

doc.add_heading('2.1 文件结构', level=2)
files_tree = [
    ('config.py', '全局配置（超参数、路径、训练设置）'),
    ('train.py', '训练主循环（含梯度累积、断点续训、码本监控）'),
    ('test_real.py', '真实物理层链路测试（BPSK+LDPC+AWGN）'),
    ('test_BPP.py', '码率统计（Bits-Per-Pixel 计算）'),
    ('models/deepsc.py', 'DeepSC 主模型（forward_train/val/test、码本利用率统计）'),
    ('models/semantic_encoder.py', '语义编码器（4层下采样，ResBlock）'),
    ('models/semantic_decoder.py', '语义解码器（4层上采样 + Skip Connection Dropout）'),
    ('models/vector_quantizer.py', 'SimVQ 量化器（冻结Embedding + 可训练投影层）'),
    ('models/channel.py', '有限码长信道模型（理论BER计算 + 比特翻转）'),
    ('losses/deepsc_loss.py', '损失函数（MSE + 加权VQ损失）'),
    ('communications/channel.py', 'AWGN / Rician 信道仿真（测试用）'),
    ('communications/modulation.py', 'BPSK / QPSK / 16-QAM 调制解调'),
    ('communications/ldpc_coding.py', 'LDPC 编译码（基于 Sionna/TensorFlow）'),
    ('communications/evaluate.py', '评估函数（test_real.py 的复用版本）'),
    ('data/datasets.py', '数据加载器（Cars196/Kodak，含预处理）'),
    ('utils/bit_utils.py', '索引↔比特流转换工具'),
    ('utils/metrics.py', 'MS-SSIM / SSIM 计算'),
]

for fname, desc in files_tree:
    p = doc.add_paragraph()
    run = p.add_run(f'{fname}')
    run.bold = True
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    p.add_run(f'  — {desc}')

doc.add_heading('2.2 数据流图（训练 vs 测试）', level=2)

doc.add_paragraph(
    '训练数据流（train.py → deepsc.forward_train）：',
)
doc.add_paragraph(
    '图像 → SemanticEncoder → [F0,F1,F2,F3] → VQ量化 → 索引 → FiniteBlocklengthChannel\n'
    '  → 理论BER比特翻转 → 损坏索引 → 还原特征 → SemanticDecoder → 重建图像\n'
    '  → MSE Loss + 加权VQ Loss（梯度通过STE直通 clean 路径）'
)

doc.add_paragraph(
    '测试数据流（test_real.py → evaluate_metrics_with_channel）：',
)
doc.add_paragraph(
    '图像 → SemanticEncoder → VQ量化 → 索引 → indices_to_bits → LDPC编码 → BPSK调制\n'
    '  → AWGN信道 → BPSK LLR软解调 → LDPC译码 → bits_to_indices → 还原特征\n'
    '  → SemanticDecoder → 重建图像 → 计算 MS-SSIM / PSNR'
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('⚠ 关键发现：训练和测试使用了完全不同的信道模型！')
run.bold = True
run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
run.font.size = Pt(12)

doc.add_page_break()

# ============================
# 三、性能问题根因诊断
# ============================
doc.add_heading('三、性能问题根因诊断', level=1)
doc.add_paragraph(
    '以下按严重程度从高到低排列，共发现 8 个影响测试性能的问题。'
    '每个问题均包含：现象、代码位置、根因分析、对性能的影响评估。'
)

# --- 问题 1 ---
doc.add_heading('3.1 CRITICAL 级别（最直接的性能杀手）', level=2)

doc.add_heading('问题 1：训练与测试的信道模型完全不同', level=3)
add_severity_badge(doc, 'CRITICAL')

doc.add_paragraph(
    '【现象】这是整个项目最致命的问题。训练时模型学习的信道噪声模式，与测试时实际经历的信道噪声模式完全不同。'
)

doc.add_paragraph('【代码位置】')
doc.add_paragraph('训练信道：models/channel.py — FiniteBlocklengthChannel.apply_channel_noise()')
add_code_block(doc, '''# 训练时的信道：理论BER公式 + 独立比特翻转
ber = self.compute_ber(snr_db, rc=rc, mod_bits=mod_bits)
# 对每个VQ索引的每个比特，以概率 ber 独立翻转
mask = torch.bernoulli(torch.full_like(bits, ber))
corrupted_bits = torch.abs(bits - mask)''')
doc.add_paragraph('测试信道：test_real.py — evaluate_metrics_with_channel()')
add_code_block(doc, '''# 测试时的信道：真实AWGN + LDPC编译码
symbols = bpsk_modulate(coded_bits_tensor)     # BPSK调制
noisy_symbols = awgn_channel(symbols, snr)      # AWGN加噪
llrs = bpsk_llr(noisy_symbols, snr, device)     # LLR软解调
decoded_bits = ldpc_decode(llrs.cpu().numpy(), ldpc_code)  # LDPC迭代译码''')

doc.add_paragraph(
    '【根因分析】两种信道模型的差异是根本性的：'
)
doc.add_paragraph(
    '1. 训练信道（理论BER模型）：假设每个比特以独立同分布的概率出错，错误均匀分布在所有时间/空间位置。'
    '这是一个"平均意义"上的信道，不存在错误突发（error burst）。'
)
doc.add_paragraph(
    '2. 测试信道（AWGN+LDPC）：AWGN 噪声在符号级加入连续值扰动，LLR 反映的是每个比特的不确定性程度。'
    'LDPC 译码存在"悬崖效应"——当 SNR 低于门限时，译码器可能输出一整块错误比特。'
    '这种突发性块错误在训练中完全没有出现。'
)
doc.add_paragraph(
    '3. 更关键的是：训练时的比特翻转直接作用于 VQ 索引的二进制表示，而测试时的噪声通过调制符号→AWGN→LLR→LDPC→硬判决，'
    '中间经历了多层非线性变换。两者的比特错误模式在统计上完全不同。'
)

doc.add_paragraph(
    '【影响评估】模型从未"见过"真实物理层链路的错误分布，无法学习对它的鲁棒性。'
    '这相当于用高斯噪声训练了一个降噪模型，却让它去处理椒盐噪声——效果必然大打折扣。'
    '这是导致 test_real.py 性能不佳的最主要原因。'
)

# --- 问题 2 ---
doc.add_heading('问题 2：梯度直通估计器使编码器无法学习信道鲁棒性', level=3)
add_severity_badge(doc, 'CRITICAL')

doc.add_paragraph('【代码位置】models/deepsc.py 第 79 行，forward_train() 方法')
add_code_block(doc, '''# 当前实现
vq_loss, quantized_clean, encoding_idx = self.vector_quantizers[i](feat)
corrupted_idx, _ = self.channel.apply_channel_noise(...)
quantized_noisy = self.vector_quantizers[i].get_quantized_features(corrupted_idx)

# ★ 问题所在：STE (Straight-Through Estimator)
quantized_final = quantized_clean + (quantized_noisy - quantized_clean).detach()
# 前向: quantized_final = quantized_noisy （解码器看到有噪特征）
# 反向: 梯度 = d(loss)/d(quantized_clean) （编码器只收到干净梯度）''')

doc.add_paragraph(
    '【根因分析】STE 的原始目的是让梯度"跳过"不可微的量化操作。但这里的使用方式产生了副作用：'
)
doc.add_paragraph(
    '1. 前向传播时，解码器接收的是 quantized_noisy（被信道噪声损坏的特征）。'
)
doc.add_paragraph(
    '2. 反向传播时，梯度通过 quantized_clean（无噪声路径）回传。'
    '.detach() 操作切断了噪声路径的梯度，编码器完全收不到"这个噪声造成了多大重建误差"的信号。'
)
doc.add_paragraph(
    '3. 结果：编码器被训练为"假设信道是完美的"来产生特征，而解码器被训练为"用有噪特征也能尽量重建"。'
    '两者优化目标不一致，编码器没有被激励去学习信道鲁棒的表征。'
)
doc.add_paragraph(
    '【影响评估】这是训练机制的架构性缺陷。即使换了更真实的信道模拟（修复问题1），'
    '如果 STE 不修正，编码器仍然无法主动产生鲁棒特征。'
)

# --- 问题 3 ---
doc.add_heading('问题 3：调制方式训练/测试不匹配', level=3)
add_severity_badge(doc, 'CRITICAL')

doc.add_paragraph('【代码位置】models/deepsc.py 第 46-52 行 _sample_mod_bits() + test_real.py 第 141 行')
add_code_block(doc, '''# 训练时：随机选取 BPSK/QPSK/16-QAM
def _sample_mod_bits(self, snr_db):
    if snr_db < 4.0:
        return random.choice([1, 2])      # BPSK 或 QPSK
    elif snr_db < 8.0:
        return random.choice([1, 2, 4])   # BPSK/QPSK/16QAM
    else:
        return random.choice([2, 4])      # QPSK 或 16QAM

# 测试时：固定使用 BPSK
symbols = bpsk_modulate(coded_bits_tensor)''')

doc.add_paragraph(
    '【根因分析】训练时模型被要求在 3 种调制方式的错误特征之间进行泛化。'
    '不同调制的 BER-SNR 曲线差异很大（BPSK 最鲁棒，16-QAM 最敏感）。'
    '测试时只使用 BPSK，但模型在训练中只有约 33% 的时间看到 BPSK 对应的错误模式。'
    '这导致模型在 BPSK 上的专项优化不足。'
)
doc.add_paragraph(
    '【影响评估】直接导致 BPSK 测试条件下的性能不是最优。修复方法很简单——让训练和测试使用相同的调制方式。'
)

doc.add_page_break()

# --- HIGH ---
doc.add_heading('3.2 HIGH 级别（显著影响性能）', level=2)

doc.add_heading('问题 4：仅使用 MSE 损失导致重建图像模糊', level=3)
add_severity_badge(doc, 'HIGH')

doc.add_paragraph('【代码位置】losses/deepsc_loss.py 第 16 行')
add_code_block(doc, '''recon_loss = self.criterion(x_hat, x)  # nn.MSELoss()''')

doc.add_paragraph(
    '【根因分析】MSE（均方误差）在数学上等价于最大化峰值信噪比（PSNR），但它有一个众所周知的缺陷：'
    '倾向于产生"所有可能解的平均"，导致重建图像模糊、缺乏高频纹理细节。'
    '而测试指标中的 MS-SSIM 衡量的是结构相似性，对模糊和纹理丢失非常敏感。'
    '用 MSE 训练却用 MS-SSIM 评估，存在天然的指标不一致。'
)
doc.add_paragraph(
    '在语义通信领域，主流做法是使用 MSE + 感知损失（LPIPS/VGG feature loss）的组合，'
    '或者直接用 MS-SSIM 作为损失函数的一部分。仅用 MSE 会导致重建图像虽然 PSNR 还可以，但 MS-SSIM 不理想。'
)

doc.add_heading('问题 5：码本容量不足 + 高维稀疏问题', level=3)
add_severity_badge(doc, 'HIGH')

doc.add_paragraph('【代码位置】config.py 第 12-13 行')
add_code_block(doc, '''EMBEDDING_DIM_LIST = [128, 256, 512, 1024]
NUM_EMBEDDINGS_LIST = [64, 64, 64, 64]  # 每层仅64个码字''')

doc.add_paragraph(
    '【根因分析】以第 3 层为例：64 个码字分布在 1024 维空间中，这是极端稀疏的。'
    '每个码字需要覆盖一个巨大的特征空间区域。在训练时还需要对抗信道噪声扰动，'
    '码本容量不足会导致：'
)
doc.add_paragraph('1. 量化误差大：编码器输出特征与最近码字的距离远，信息损失大。')
doc.add_paragraph('2. 码本坍缩风险：部分码字从未被使用（死码字），实际有效码字数更少。')
doc.add_paragraph(
    '3. SimVQ 的冻结 Embedding 机制进一步限制了码本自适应能力——只有投影层可训练，'
    '底层码字位置无法调整（只通过投影矩阵做线性变换）。'
)
doc.add_paragraph(
    '建议：关注训练日志中每 10 epoch 的"码本利用率统计报告"，如果 active_ratio < 80% 或'
    ' collapse_ratio > 10%，说明码本坍缩严重。'
)

doc.add_heading('问题 6：训练/测试图像分辨率和数据域不匹配', level=3)
add_severity_badge(doc, 'HIGH')

doc.add_paragraph('【代码位置】data/datasets.py 第 42-62 行 + config.py 第 44-46 行')
add_code_block(doc, '''# 训练：Cars196 → Resize(256) → RandomCrop(256)
# 测试：Kodak → Resize((768, 512))  ← 不同分辨率、不同数据集''')

doc.add_paragraph(
    '【根因分析】存在两个层面的分布偏移（domain shift）：'
)
doc.add_paragraph(
    '1. 数据域偏移：Cars196 全是汽车图像（背景相对简单），Kodak 是通用自然图像（含复杂纹理、人像、风景等）。'
    'VQ 码本在汽车特征上训练，遇到 Kodak 的多样化纹理时覆盖不足。'
)
doc.add_paragraph(
    '2. 分辨率偏移：训练时特征图的空间分布是 128×128, 64×64, 32×32, 16×16（基于 256 输入），'
    '测试时变成了 384×256, 192×128, 96×64, 48×32（基于 768×512 输入）。'
    '虽然全卷积网络可以处理任意分辨率，但 VQ 码本是在特定分辨率下的特征分布上训练的，'
    '不同分辨率下特征的均值和方差可能不同，导致量化效果下降。'
)

doc.add_page_break()

# --- MEDIUM ---
doc.add_heading('3.3 MEDIUM 级别', level=2)

doc.add_heading('问题 7：VQ 损失权重失衡', level=3)
add_severity_badge(doc, 'MEDIUM')

doc.add_paragraph('【代码位置】config.py 第 17 行')
add_code_block(doc, '''LAYER_LOSS_WEIGHTS = [1, 1, 5, 10]  # Layer3权重是Layer0的10倍''')

doc.add_paragraph(
    '【根因分析】深层权重高的设计意图是合理的——深层包含更多语义信息。但 10 倍的悬殊差距带来两个问题：'
)
doc.add_paragraph(
    '1. 训练被 Layer 3 的 VQ 损失主导，浅层（Layer 0/1）的量化质量被忽视。'
    '而浅层包含大量纹理和边缘细节，对 MS-SSIM 和视觉质量至关重要。'
)
doc.add_paragraph(
    '2. 不同层的 VQ 损失量级本就不同（embedding_dim 不同），额外乘权重加剧了不平衡。'
)

doc.add_heading('问题 8：跳跃连接 Dropout 在训练和推理间存在 gap', level=3)
add_severity_badge(doc, 'MEDIUM')

doc.add_paragraph('【代码位置】config.py 第 21 行 + models/semantic_decoder.py 第 6-21 行')
add_code_block(doc, '''SKIP_DROPOUT_P = [0.5, 0.45, 0.05]  # Layer0浅层丢弃50%
# 但在推理时：model.eval() → 所有skip 100%通过''')

doc.add_paragraph(
    '【根因分析】训练时 Layer 0 的跳跃连接有 50% 概率被整条丢弃，解码器学会了"不依赖 skip 也能重建"。'
    '但推理时所有 skip 全部激活，突然的信息增量可能反而干扰解码器的重建逻辑。'
    '这种 train/test 行为不一致可能对性能有微弱负面影响。'
)

doc.add_page_break()

# ============================
# 四、优化方案
# ============================
doc.add_heading('四、优化方案（3套方案对比）', level=1)

doc.add_paragraph(
    '针对以上 8 个问题，提出三套递进式方案。方案 A 最小改动，方案 B 治本为主，方案 C 全面重构。'
)

doc.add_heading('4.1 方案对比总览', level=2)

add_table_simple(doc,
    ['对比维度', '方案 A（快速修复）', '方案 B（治本增强）', '方案 C（全面重构）'],
    [
        ['改动量', '约 30 行', '约 100 行', '约 300+ 行'],
        ['解决问题', '问题1(部分)、3、7', '问题1~8（除8外全部）', '全部8个问题'],
        ['需要重新训练', '否（可finetune）', '是', '是'],
        ['预期MS-SSIM提升', '0.02-0.05', '0.05-0.10', '0.08-0.15'],
        ['风险', '低', '中', '高'],
        ['实施周期', '1-2天', '3-5天', '1-2周'],
    ]
)

doc.add_heading('4.2 方案 A：最小改动快速验证', level=2)

doc.add_paragraph('改动 1：修复训练信道模拟（约 20 行）')
doc.add_paragraph(
    '在 forward_train 中，将 FiniteBlocklengthChannel 的理论 BER 方式替换为：'
    '索引 → 二进制比特 → BPSK 调制 → 加 AWGN 噪声 → 硬判决 → 还原索引。'
    '保持 STE 机制不变，但让噪声模式更接近真实链路。'
)

doc.add_paragraph('改动 2：固定调制为 BPSK（约 3 行）')
doc.add_paragraph(
    '将 _sample_mod_bits() 简化为始终返回 1，与测试对齐。'
)

doc.add_paragraph('改动 3：均衡 VQ 损失权重（约 1 行）')
doc.add_paragraph(
    '将 LAYER_LOSS_WEIGHTS 从 [1, 1, 5, 10] 改为 [3, 3, 3, 3] 或 [2, 2, 3, 3]。'
)

doc.add_paragraph('优点：改动极小，可快速验证方向；风险低；可在现有 checkpoint 上微调。')
doc.add_paragraph('缺点：未解决问题2（STE梯度）、问题4（MSE损失）、问题5（码本容量）等深层问题。')

doc.add_heading('4.3 方案 B：治本修复 + 适当增强', level=2)

doc.add_paragraph('在方案 A 基础上，额外增加：')

doc.add_paragraph('改动 4：可微分 AWGN 信道 + Gumbel-Softmax 软量化（约 50 行）')
doc.add_paragraph(
    '替换 .detach() STE 模式：使用 Gumbel-Softmax 或软量化（soft assignment），'
    '让信道噪声的梯度能回传到编码器。具体方案：'
)
doc.add_paragraph(
    '• 索引 → one-hot → 软分配（softmax over codebook distances with temperature τ）→ 软特征 → 解码器'
)
doc.add_paragraph(
    '• 训练时在软特征上加入 AWGN 噪声（完全可微），梯度可以端到端流过编码器和解码器'
)
doc.add_paragraph(
    '• 训练初期 τ 较大（软分配接近均匀），逐渐退火到 τ 较小（接近硬分配）'
)
doc.add_paragraph(
    '这样编码器就能直接感知"信道噪声如何影响重建质量"，主动学习鲁棒表征。'
)

doc.add_paragraph('改动 5：增加感知损失（约 30 行）')
doc.add_paragraph(
    '在 DeepSCLoss 中增加 LPIPS（Learned Perceptual Image Patch Similarity）损失：'
    'total_loss = MSE_loss + 0.1 * LPIPS_loss + weighted_VQ_loss'
)
doc.add_paragraph(
    'LPIPS 基于预训练 VGG/ AlexNet 特征，比 MSE 更贴近人类视觉感知，能直接提升 MS-SSIM。'
)

doc.add_paragraph('改动 6：扩大码本容量（约 1 行）')
doc.add_paragraph(
    '将 NUM_EMBEDDINGS_LIST 从 [64, 64, 64, 64] 提升到 [128, 256, 256, 256] 或 [256, 256, 256, 512]，'
    '缓解码本稀疏问题。注意：码本增大会增加比特率，需要在 BPP 和重建质量之间权衡。'
)

doc.add_paragraph('优点：根本性解决训练测试不一致 + 梯度问题 + 损失单一；预期效果最好。')
doc.add_paragraph('缺点：需要重新训练；Gumbel-Softmax 训练初期可能不稳定，需要温度退火调度。')

doc.add_heading('4.4 方案 C：全面重构', level=2)

doc.add_paragraph('在方案 B 基础上额外增加：')

doc.add_paragraph('改动 7：训练时加入 LDPC 等效模拟')
doc.add_paragraph(
    '由于 LDPC（Sionna/TensorFlow）不可微，可以在训练时使用 SNR 相关的等效 BER 曲线来近似 LDPC 编译码效果，'
    '或者使用可微的卷积码/Turbo码近似替代。'
)

doc.add_paragraph('改动 8：多分辨率训练')
doc.add_paragraph(
    '训练时随机使用 [256×256, 384×256, 512×384, 768×512] 多种分辨率，使模型适应不同输入尺寸。'
)

doc.add_paragraph('改动 9：码本 K-Means 初始化 + EMA 更新')
doc.add_paragraph(
    '用编码器在训练数据上提取的特征做 K-Means 聚类来初始化码本。'
    '加入 EMA（指数移动平均）更新码字，替代完全冻结的 Embedding。'
)

doc.add_paragraph('改动 10：GAN 判别器增强')
doc.add_paragraph('在解码器后增加 PatchGAN 判别器，用对抗训练提升纹理真实感。')

doc.add_paragraph('优点：理论最优，全面解决所有问题。')
doc.add_paragraph('缺点：工程量大，训练不稳定风险高（GAN + VQ + 软量化 三重不稳定因素叠加）。')

doc.add_page_break()

# ============================
# 五、推荐方案详细说明
# ============================
doc.add_heading('五、推荐方案：方案 B 详细说明', level=1)

doc.add_paragraph(
    '综合考虑 MS-SSIM > 0.85 的当前基线、训练已接近完成（>200 epochs）的情况，'
    '以及"最小改动快速验证 + 治本"的需求，推荐方案 B。以下详细展开各项改动的具体实现思路。'
)

doc.add_heading('5.1 改动 1：可微分 AWGN 信道训练', level=2)

doc.add_paragraph('目标：让训练时的信道噪声模式与测试时一致，且梯度可穿过信道到达编码器。')

doc.add_paragraph('实现思路：')
doc.add_paragraph(
    '1. 在 forward_train() 中，将索引转为 one-hot 或使用编码器输出特征与码本的距离，'
    '通过 softmax 得到"软分配权重"，乘以码本得到软量化特征。'
)
doc.add_paragraph(
    '2. 在软量化特征上直接加入 AWGN 噪声（使用 torch.randn_like），噪声功率根据 SNR 计算。'
)
doc.add_paragraph(
    '3. 解码器接收加噪的软量化特征，计算重建损失。梯度端到端流过编码器→软量化→加噪→解码器。'
)
doc.add_paragraph(
    '4. 同时保留原有的硬量化 + VQ 损失（commitment loss）作为辅助损失，确保码本正常训练。'
)
doc.add_paragraph(
    '5. 推理时（forward_test）仍然使用硬量化 + 真实物理层链路，不受影响。'
)

doc.add_heading('5.2 改动 2：固定调制为 BPSK', level=2)
doc.add_paragraph(
    '删除 _sample_mod_bits() 的随机选择逻辑，固定返回 1（BPSK），'
    '使训练和测试使用完全一致的调制方式。如果后续需要支持多种调制，'
    '可以改为在训练时也按比例混合，但测试时需要分别评估。'
)

doc.add_heading('5.3 改动 3：均衡 VQ 损失权重', level=2)
doc.add_paragraph(
    '将 LAYER_LOSS_WEIGHTS 从 [1, 1, 5, 10] 改为 [2, 2, 3, 3]，'
    '适当增加浅层权重以保护纹理细节，同时深层仍有适度优势以保持语义质量。'
)

doc.add_heading('5.4 改动 4：增加感知损失（LPIPS）', level=2)
doc.add_paragraph(
    'LPIPS 使用预训练的 AlexNet 或 VGG 网络提取多层特征，计算特征空间的 L2 距离。'
    '与 MS-SSIM 高度相关，是被广泛验证的感知度量。'
)
doc.add_paragraph('混合损失公式：')
add_code_block(doc, '''total_recon_loss = MSE(x, x_hat) + λ_lpips * LPIPS(x, x_hat)
# λ_lpips 建议从 0.05 开始，可视情况调整''')
doc.add_paragraph(
    '注意：LPIPS 要求输入在 [-1, 1] 范围（与当前 normalize(0.5, 0.5) 一致），无需修改预处理。'
)

doc.add_heading('5.5 改动 5：扩大码本容量', level=2)
doc.add_paragraph(
    '将码本大小从 [64, 64, 64, 64] 扩大到 [128, 256, 256, 256]。'
    '码本增大带来的额外比特率增加（每 token 从 6 bit 到 7 bit（Layer 0）和 8 bit（Layer 1-3）），'
    '可以通过降低 CHANNEL_CODING_RATE 或调整 BLOCK_LENGTH 来补偿。'
)

doc.add_heading('5.6 预期效果', level=2)

doc.add_paragraph('各改动对性能指标的预期贡献：')
add_table_simple(doc,
    ['改动', '预期 MS-SSIM 提升', '预期 PSNR 提升', '置信度'],
    [
        ['可微分AWGN信道训练', '+0.03 ~ 0.05', '+1.0 ~ 2.0 dB', '高'],
        ['固定BPSK调制', '+0.01 ~ 0.02', '+0.3 ~ 0.5 dB', '高'],
        ['均衡VQ权重', '+0.01 ~ 0.02', '+0.2 ~ 0.5 dB', '中'],
        ['LPIPS感知损失', '+0.02 ~ 0.04', '-0.5 ~ +0.5 dB', '高（MS-SSIM提升显著，PSNR可能略降）'],
        ['扩大码本容量', '+0.02 ~ 0.04', '+0.5 ~ 1.5 dB', '中高'],
        ['合计', '+0.09 ~ 0.17', '+1.5 ~ 5.0 dB', '-'],
    ]
)

doc.add_page_break()

# ============================
# 六、总结与建议
# ============================
doc.add_heading('六、总结与建议', level=1)

doc.add_heading('6.1 核心结论', level=2)
doc.add_paragraph(
    'test_real.py 性能不佳的根本原因是训练和测试之间存在系统性的"信道模型鸿沟"（Channel Model Gap）。'
    '训练使用理论 BER + 独立比特翻转 + STE 梯度截断，测试使用 AWGN + LDPC + 真实调制解调。'
    '这两个信道的行为模式在统计上有本质差异，模型不可能在未见过的信道条件下表现良好。'
)

doc.add_heading('6.2 推荐实施路线', level=2)

doc.add_paragraph(
    '第一阶段（1-2天，方案A的改动1-3）：快速验证信道一致性的重要性。'
)
doc.add_paragraph('• 在现有 checkpoint 上进行微调（finetune），而非从头训练')
doc.add_paragraph('• 如果 test_real.py 指标有明显提升（>0.02 MS-SSIM），说明方向正确')
doc.add_paragraph('• 使用较小的学习率（如 1e-6），微调 10-20 个 epoch')

doc.add_paragraph(
    '第二阶段（3-5天，方案A确认有效后实施方案B全部改动）：'
)
doc.add_paragraph('• 从头开始训练（因为模型架构和损失函数有较大变化）')
doc.add_paragraph('• 密切关注码本利用率统计，确保码本不坍缩')
doc.add_paragraph('• 监控 LPIPS 损失和 MSE 损失的平衡')

doc.add_paragraph(
    '第三阶段（可选，方案C的改动9-10）：如果方案B的效果仍不满足需求，再追加强化。'
)

doc.add_heading('6.3 关键风险提示', level=2)
doc.add_paragraph(
    '1. Gumbel-Softmax / 软量化训练可能初期不稳定，需要仔细调温度退火策略。'
)
doc.add_paragraph(
    '2. 增大码本会增加 BPP（比特率），需要在重建质量和传输效率之间权衡。'
)
doc.add_paragraph(
    '3. LPIPS 需要额外显存（约 200MB for AlexNet），确认 GPU 显存充足。'
)
doc.add_paragraph(
    '4. 方案B需要重新训练，建议保留当前最佳 checkpoint 作为 baseline 对比。'
)

# ===== 保存 =====
doc.save(OUTPUT_PATH)
print(f"报告已生成: {OUTPUT_PATH}")
