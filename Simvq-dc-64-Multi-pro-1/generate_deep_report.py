"""
生成 SimVQ 项目深度性能分析报告（第二版）
包含: BN vs GN 深入对比, SimVQ 可行性数学分析, 快速评估方法论, 以及逐层逐模块的微观分析
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "SimVQ_深度性能分析报告_v2.docx")

doc = Document()

# ===== 全局样式 =====
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.25

for level in range(1, 5):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '微软雅黑'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if level == 1:
        hs.font.size = Pt(18); hs.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    elif level == 2:
        hs.font.size = Pt(14); hs.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    elif level == 3:
        hs.font.size = Pt(12); hs.font.color.rgb = RGBColor(0x34, 0x49, 0x5E)
    else:
        hs.font.size = Pt(11); hs.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

def add_code(doc, code, font_size=8.5):
    """添加代码块"""
    for line in code.strip().split('\n'):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.line_spacing = 1.0

def add_bold_text(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r+1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

def add_warning_box(doc, title, text):
    p = doc.add_paragraph()
    run = p.add_run(f'⚠ {title}：')
    run.bold = True; run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C); run.font.size = Pt(11)
    run2 = p.add_run(text)
    run2.font.size = Pt(10.5)

def add_info_box(doc, title, text):
    p = doc.add_paragraph()
    run = p.add_run(f'ℹ {title}：')
    run.bold = True; run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB); run.font.size = Pt(11)
    run2 = p.add_run(text)
    run2.font.size = Pt(10.5)

# ============= 封面 =============
doc.add_paragraph()
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('SimVQ 语义通信系统\n深度性能诊断报告')
run.bold = True; run.font.size = Pt(26); run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

doc.add_paragraph()
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('—— 架构级根因分析与快速验证方法论 ——')
run.font.size = Pt(13); run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

doc.add_paragraph()
dt = doc.add_paragraph(); dt.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = dt.add_run('版本 V2.0  |  2026年5月14日')
run.font.size = Pt(11); run.font.color.rgb = RGBColor(0x95, 0xA5, 0xA6)

doc.add_page_break()

# ============= 目录 =============
doc.add_heading('报告结构', level=1)
toc = [
    '第一部分：宏观架构审查 —— 训练/测试"两层皮"问题的微观解剖',
    '  1.1 双信道模型的逐层对比',
    '  1.2 FiniteBlocklengthChannel 的 5 个理论缺陷',
    '  1.3 梯度直通估计器 (STE) 的完整梯度流分析',
    '第二部分：Normalization 层选型分析 —— BN vs GN vs IN',
    '  2.1 项目中 BN 的完整分布图 (20个BN层盘点)',
    '  2.2 BN 在信道语义通信场景下的 6 个失效模式',
    '  2.3 GN 的数学优势与适配性论证',
    '第三部分：SimVQ 模块深度解析',
    '  3.1 SimVQ 与传统 VQ-VAE 的对比',
    '  3.2 冻结嵌入 + 投影层的数学局限性',
    '  3.3 K<<D 高维稀疏码本的几何分析',
    '  3.4 码本坍缩风险的深层机制',
    '第四部分：逐模块微观问题清单',
    '  4.1 编码器设计问题',
    '  4.2 解码器设计问题',
    '  4.3 信道模块设计问题',
    '  4.4 损失函数设计问题',
    '  4.5 数据处理问题',
    '第五部分：快速验证方法论 (10-50 Epoch判断)',
    '  5.1 代理指标体系 (6个指标)',
    '  5.2 实用评估协议',
    '  5.3 各指标的早期预警阈值',
]
for item in toc:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.1

doc.add_page_break()

# ================================================================
# 第一部分
# ================================================================
doc.add_heading('第一部分：宏观架构审查 —— 训练/测试"两层皮"问题的微观解剖', level=1)

doc.add_paragraph(
    '这是本项目最根本的结构性问题。训练（train.py → deepsc.forward_train）和测试（test_real.py → evaluate_metrics_with_channel）'
    '使用了两个完全不同的物理层传输链路。以下从信号处理的每一步进行逐层解剖。'
)

doc.add_heading('1.1 双信道模型的逐层对比', level=2)

doc.add_paragraph('下表对标了训练和测试链路中每一层的信号处理操作，揭示了它们的分歧点：')

add_table(doc,
    ['处理阶段', '训练链路 (forward_train)', '测试链路 (test_real.py)', '是否一致?', '影响'],
    [
        ['① 信源编码', 'Encoder → 多尺度特征', 'Encoder → 多尺度特征', '✓ 一致', '无'],
        ['② VQ量化', '查码本→量化特征+索引', '查码本→索引(仅返回索引)', '✓ 一致', '无'],
        ['③ 比特表示', '索引直接转二进制(隐式)', '索引→indices_to_bits→显式bit流', '≈ 类似', '低'],
        ['④ 信道编码', '无(FiniteBlocklengthChannel\n模拟BER)', 'LDPC编码(k=128→n=256)', '✗ 完全不同', '极高'],
        ['⑤ 调制', '无(mod_bits参数影响BER,\n但不实际调制)', 'BPSK调制(0→-1, 1→+1)', '✗ 完全不同', '极高'],
        ['⑥ 信道传输', '理论BER做逐比特\n独立翻转(Bernoulli)', 'AWGN加连续噪声\n(y=x+n, n~CN(0,σ²))', '✗ 完全不同', '极高'],
        ['⑦ 接收处理', '无(直接得到损坏索引)', 'LLR软解调+BP迭代译码', '✗ 完全不同', '极高'],
        ['⑧ 特征还原', 'codebook[corrupted_idx]\n→损坏的量化特征', 'codebook[recovered_idx]\n→干净的量化特征', '≈ 类似', '中'],
        ['⑨ 解码重建', 'Decoder(损坏特征)→重建图', 'Decoder(干净特征)→重建图', '✗ 不同!', '高'],
    ],
    col_widths=[2.2, 3.5, 3.5, 1.5, 1.2]
)

doc.add_paragraph(
    '关键发现：9个处理阶段中有 5 个完全不同，1 个部分不同。这意味着模型在训练中"以为"的信道行为，'
    '与测试时实际经历的信道行为，在超过一半的处理步骤上存在差异。'
)

doc.add_heading('1.2 FiniteBlocklengthChannel 的 5 个理论缺陷', level=2)

doc.add_paragraph(
    '训练使用的 FiniteBlocklengthChannel（models/channel.py）基于 Polyanskiy 等人的有限码长信息论公式。'
    '该公式本身是正确的，但在此项目中的应用方式存在 5 个问题：'
)

doc.add_heading('缺陷 1：BER 公式与 LDPC 实际性能的系统性偏差', level=3)
doc.add_paragraph(
    'FiniteBlocklengthChannel 使用正态近似（Normal Approximation）计算给定块长和速率下的平均错误概率：'
)
add_code(doc, '''C  = log₂(1 + γ)                          # 香农容量 (AWGN, 连续输入)
V  = (1 - (1+γ)⁻²) × (log₂e)²              # 信道色散 (Channel Dispersion)
ρ  = Q(√L × (C - R) / √V)                   # 正态近似错误概率
BER = ρ / k_info_bits                        # 近似误比特率''')

doc.add_paragraph(
    '问题在于：正态近似给出的是"所有可能编码方案在最大似然译码下的平均性能上界"。'
    '但实际使用的 LDPC + BP（置信传播）译码的性能会偏离这个平均值：'
)
doc.add_paragraph('• 在低 SNR 区域，BP 译码可能比 ML 译码差 1-3 dB')
doc.add_paragraph('• 在高 SNR 区域，LDPC 存在"错误平层"（Error Floor），正态近似无法捕捉')
doc.add_paragraph('• 正态近似假设连续输入分布，但实际使用的是离散星座（BPSK/QPSK）')

doc.add_heading('缺陷 2：用平均 BER 掩盖了错误的突发性', level=3)
doc.add_paragraph(
    '正态近似给出一个标量 BER 值。随后该 BER 被当作独立伯努利概率，对每个 VQ 索引比特做独立翻转：'
)
add_code(doc, '''mask = torch.bernoulli(torch.full_like(bits, ber))   # 逐比特独立翻转
corrupted_bits = torch.abs(bits - mask)''')

doc.add_paragraph(
    '但实际 LDPC 译码的错误模式是完全不同的：当 SNR 低于译码门限时，整个 LDPC 块（256 个编码比特 = 128 个信息比特）'
    '可能同时出错。这是一个"块级"突发错误。用独立伯努利翻转来训练，模型学会的是"均匀分散的像素级噪声鲁棒性"，'
    '而测试时面对的是"集中的块级破坏"。这两种错误模式对图像重建的影响完全不同。'
)

doc.add_heading('缺陷 3：Block_Length 参数的不一致使用', level=3)
doc.add_paragraph(
    'FiniteBlocklengthChannel 被配置为 BLOCK_LENGTH = 256（coded_block_length_bits），即理论公式按 256 比特的块来计算 BER。'
    '但实际传输的 VQ 索引比特流长度约为 130,560 比特（所有层索引拼接后）。'
    '有限码长理论的核心思想是"短块长导致更高的错误概率"，但 130K 比特已经不短了 —— '
    '实际 LDPC 方案设 n=256 也是一个短码。两者的"短"含义不同：理论公式的 256 是单个码字的块长，'
    '而实际有 ~1000 个这样的码字并行传输。理论公式用一个块来分析，却把结果当作独立事件应用到所有比特，这在数学上是不严格的。'
)

doc.add_heading('缺陷 4：mod_bits 只是 BER 公式的输入参数，不产生真实调制效应', level=3)
doc.add_paragraph(
    '训练时 _sample_mod_bits() 在 {1,2,4} 中随机选择，但 mod_bits 仅仅改变了 BER 计算公式中的 R_transport = rc × mod_bits。'
    '它不产生真实的星座映射、符号间干扰、或 I/Q 不平衡等实际调制会引入的效应。'
    '模型接收到的梯度信号是"当前 SNR 条件下有一个大概的 BER"，没有学到任何调制相关的特征。'
)

doc.add_heading('缺陷 5：Clamping 引入的码字偏差', level=3)
doc.add_paragraph(
    '当比特翻转导致索引值超出合法范围时，代码使用 clamp 处理：'
)
add_code(doc, '''corrupted_indices = torch.clamp(corrupted_indices, 0, num_embeddings - 1)''')
doc.add_paragraph(
    '例如，Layer 3 有 64 个码字（索引 0-63，6 比特表示）。当 BER 较高时，比特翻转可能导致索引值 > 63。'
    '这些溢出值全部被 clamp 到 63。这意味着在高 BER（低 SNR）条件下，索引 63 被人为地"过采样"了。'
    '解码器在训练时会过于频繁地看到码字 63 的特征，形成偏差。这个偏差在测试时不存在（LDPC 译码不会产生溢出索引），'
    '但训练时解码器学到的"码字 63 很常见"的先验仍然保留在权重中。'
)

doc.add_heading('1.3 梯度直通估计器 (STE) 的完整梯度流分析', level=2)

doc.add_paragraph(
    '现在对 forward_train 中的关键行（deepsc.py:79）进行完整的梯度流追踪，揭示为什么编码器学不到信道鲁棒性。'
)

doc.add_heading('前向传播路径', level=3)
add_code(doc, '''# Step 1: 编码器产生特征
feat = encoder(x)                                    # 有梯度

# Step 2: VQ 量化 (STE 第一次)
vq_loss, quantized_clean, idx = vq(feat)             # quantized_clean 有梯度(STE通过argmin)

# Step 3: 信道损坏索引
corrupted_idx = channel.apply_noise(idx)              # 无梯度 (非PyTorch操作/随机操作)

# Step 4: 从损坏索引查码本
quantized_noisy = vq.get_quantized_features(corrupted_idx) # 无梯度 (索引是整数,不可微)

# Step 5: 组合 (STE 第二次)
quantized_final = quantized_clean + (quantized_noisy - quantized_clean).detach()
# = quantized_noisy (前向)  ← 解码器实际看到的是有噪特征

# Step 6: 解码
recon = decoder(quantized_final)                      # 基于有噪特征重建

# Step 7: 损失
loss = MSE(recon, x) + weighted_vq_loss               # 基于有噪重建计算误差''')

doc.add_heading('反向传播路径 (梯度链)', level=3)
add_code(doc, '''∂L/∂(quantized_final)  →  解码器参数  (∂L/∂decoder_params)

# 关键分歧点：
∂(quantized_final)/∂(quantized_clean) = 1            # 梯度通过clean路径
∂(quantized_final)/∂(quantized_noisy) = 0            # .detach() 截断了噪声路径！

# 因此：
∂L/∂(quantized_clean) = ∂L/∂(quantized_final) × 1   # 编码器收到梯度
# 但这个梯度是基于"clean特征 → 解码 → 计算loss"的假设！
# 实际loss是基于 noisy特征计算的，但梯度却指向clean特征方向！

# 更深层：
∂L/∂(encoder_params) = ∂L/∂(quantized_clean) × ∂(quantized_clean)/∂feat × ∂feat/∂(encoder_params)
# = ∂L/∂(quantized_final) × ∂(quantized_clean)/∂feat × ∂feat/∂(encoder_params)
#                     ↑ 这个梯度说："调整编码器使clean特征产生更好的重建"
#                     但实际上前向用的是noisy特征！梯度方向存在系统性偏差！''')

doc.add_paragraph(
    '这意味着编码器被优化的方向是"如何让干净特征通过解码器重建得更好"，'
    '而不是"如何让编码器特征在信道损坏后仍然能被解码器有效利用"。'
    '编码器的参数更新完全不包含信道噪声的影响信息。'
)

doc.add_paragraph(
    '用一个具体例子说明：假设在某个训练步骤，SNR=3dB，码字A被错误地翻转为码字B。'
    '解码器收到了码字B的特征，产生了很差的recon，loss很高。梯度反向传播时：'
)
doc.add_paragraph(
    '• 解码器学到了："当收到码字B的特征时，要调整权重以减少重建误差"（这是正确的学习信号）'
)
doc.add_paragraph(
    '• 编码器收到了 grad = ∂L/∂quantized_clean，但这个梯度会推动编码器参数使 quantized_clean 发生变化。'
    '然而 quantized_clean 是来自正确码字A的特征，编码器被推动去"改善码字A的特征"，而不是"让特征对误码B更鲁棒"。'
    '这完全搞错了学习方向！'
)
doc.add_paragraph(
    '正确的梯度应该告诉编码器："你的特征和码字B太接近了，请把它们拉远一些，'
    '或者让码字A和码字B的特征在解码器看来更相似（以减少误码影响）"。'
    '但当前的STE机制传递的是错误方向的信号。'
)

doc.add_page_break()

# ================================================================
# 第二部分
# ================================================================
doc.add_heading('第二部分：Normalization 层选型分析', level=1)

doc.add_heading('2.1 项目中 BN 的完整分布图', level=2)

doc.add_paragraph('全模型共包含 20 个 BatchNorm2d 层，分布在以下模块中：')

add_table(doc,
    ['模块', '包含的 BN 数量', '输入特征图大小(训练)', 'BN 统计样本量/次'],
    [
        ['SemanticEncoder (4×DownSampleBlock)', '4×3 = 12', '256→128→64→32→16', '24×(H×W)'],
        ['SemanticDecoder (4×UpSampleBlock)', '4×2 = 8', '16→32→64→128→256', '24×(H×W)'],
        ['合计', '20', '-', '-'],
    ],
    col_widths=[5, 2.5, 4, 4]
)

doc.add_paragraph(
    '每个 BN 层维护两个缓冲区：running_mean 和 running_var（指数移动平均），'
    '在 model.train() 时用 batch 统计更新，在 model.eval() 时使用这些 running 统计。'
)

doc.add_heading('2.2 BN 在信道语义通信场景下的 6 个失效模式', level=2)

doc.add_heading('失效模式 1：SNR 混合导致的 running statistics 失真', level=3)
doc.add_paragraph(
    '训练时 SNR 从 [0, 15] dB 均匀采样。不同 SNR 下，经过 VQ+信道后的特征分布完全不同：'
)
doc.add_paragraph('• SNR=0dB：比特错误率极高（BER≈0.15），损坏特征来自随机码字 → 特征方差极大')
doc.add_paragraph('• SNR=15dB：比特错误率极低（BER≈1e-7），损坏特征≈干净特征 → 特征方差较小')
doc.add_paragraph(
    'BN 的 running_mean 和 running_var 是对所有 SNR 条件下 batch 统计的指数移动平均。'
    '这相当于把 0dB 的高方差和 15dB 的低方差"混合"成一个平均值。测试时，'
    '每个 SNR 是单独评估的，但 BN 使用的 running statistics 却是"平均 SNR"的统计量。'
)
doc.add_paragraph(
    '具体来说：测试 SNR=0dB 时，解码器特征的实际方差远大于 running_var（因为 running_var 被高 SNR 拉低了），'
    '导致 BN 的归一化不足（除数偏小），激活值偏大。测试 SNR=12dB 时则相反，特征方差小于 running_var，'
    'BN 过度归一化（除数偏大），信号被压缩。'
)
doc.add_paragraph(
    '这个"归一化偏差"在所有 20 个 BN 层中累积传播，对解码器的重建质量产生显著影响。'
)

doc.add_heading('失效模式 2：训练/推理时 BN 行为不一致', level=3)
doc.add_paragraph(
    '训练时 (model.train())：BN 使用当前 mini-batch 的统计量进行归一化（μ_batch, σ²_batch）。'
    '这些统计量反映了训练信道（独立比特翻转）的特征分布。'
)
doc.add_paragraph(
    '推理时 (model.eval())：BN 使用 running_mean 和 running_var。'
    '但测试时的信道是 AWGN+LDPC，其产生的特征分布与训练时不同。'
    '即使 running statistics 完美地表示了训练分布，它们在测试分布上也可能是错误的。'
)
doc.add_paragraph(
    '这进一步加剧了训练/测试不一致问题——不仅信道模型不同，连 normalization 层的参数也存在偏差。'
)

doc.add_heading('失效模式 3：batch_size 对统计估计的影响', level=3)
doc.add_paragraph(
    '当前 MICRO_BATCH_SIZE = 24，TOTAL_BATCH_SIZE = 24，accumulation_steps = 1。'
    'BN 仅在 24 个样本上估计均值和方差。对于 16×16×1024 的特征图（Layer 3），'
    '每个 BN 在 24×16×16 = 6144 个值上计算统计量。虽然 6144 个样本点足够估计均值和方差，'
    '但 24 个独立样本的代表性有限——如果 batch 中碰巧包含 2-3 张异常图像，BN 统计就会偏离。'
)
doc.add_paragraph(
    '实际上代码中有一个 BN momentum 调整逻辑（train.py:62-68），但因为 accumulation_steps=1 而从未触发。'
    '这说明代码作者意识到了 BN 在梯度累积场景下的问题，但当前配置下没有实际效果。'
)

doc.add_heading('失效模式 4：跨分辨率泛化问题', level=3)
doc.add_paragraph(
    '训练时输入为 256×256，测试时输入为 768×512。虽然卷积 + BN 理论上可以处理任意分辨率，'
    '但 BN 的 running statistics 是在特定分辨率下累积的。即使特征的空间尺寸不同，'
    'BN 假设每个通道的均值和方差不随空间维度变化。这个假设在以下情况下可能不成立：'
)
doc.add_paragraph(
    '• 不同分辨率下，卷积核的感受野覆盖了不同的图像内容比例'
)
doc.add_paragraph(
    '• 编码器输出特征的统计量可能因输入分辨率不同而变化（更大的图→更多的空间上下文→不同的特征分布）'
)
doc.add_paragraph(
    '特别地，对于测试的 768×512 输入，编码器各层输出尺寸为 (384×256, 192×128, 96×64, 48×32)，'
    '而训练时为 (128×128, 64×64, 32×32, 16×16)。虽然每个位置的感受野大小相同，但全局统计量可能不同。'
)

doc.add_heading('失效模式 5：编码器和解码器中 BN 的非对称问题', level=3)
doc.add_paragraph(
    '编码器的 BN（12 层）处理的是干净的原始图像特征——它们的统计量相对稳定，'
    '只取决于图像内容的变化。解码器的 BN（8 层）处理的是经过 VQ 量化 + 信道损坏的特征——'
    '统计量高度依赖 SNR 和信道错误模式。'
)
doc.add_paragraph(
    '解码器 BN 的 running statistics 需要同时适应：(a) 不同 SNR 的噪声水平，(b) 不同码字选择的特征变化，'
    '(c) 不同图像内容的特征分布。这是一个极其复杂的分布混合，一个简单的 running mean/var 几乎不可能准确表示。'
)

doc.add_heading('失效模式 6：与 PReLU 的交互问题', level=3)
doc.add_paragraph(
    '每个 BN 后面紧跟 PReLU 激活（learnable negative slope）。PReLU 的可学习参数对输入分布敏感。'
    '当 BN 的归一化因 SNR 变化而偏移时（失效模式 1），PReLU 的有效行为也会改变：'
)
doc.add_paragraph(
    '• 过度归一化（实际方差 < running_var）→ 激活值被压缩 → PReLU 工作在线性区域，表达能力减弱'
)
doc.add_paragraph(
    '• 归一化不足（实际方差 > running_var）→ 激活值被放大 → PReLU 的负斜率部分被过度使用，可能产生大量负值'
)
doc.add_paragraph(
    '由于 PReLU 的参数在训练中已经学到了适应"平均 SNR"的分布，测试时特定 SNR 的偏移会导致次优的激活模式。'
)

doc.add_heading('2.3 GN 的数学优势与适配性论证', level=2)

doc.add_heading('Group Normalization 的工作原理', level=3)
doc.add_paragraph(
    'GN 将通道分为 G 组，在每组内沿 (H, W) 维度计算均值和方差，进行归一化。'
    '关键区别：GN 不依赖 batch 维度，每个样本独立归一化。'
)
add_code(doc, '''# BN: 对 (N, H, W) 归一化 → 依赖 batch 中其他样本
y_bn = (x - μ_batch) / √(σ²_batch + ε) × γ + β   # μ,σ² 跨样本计算

# GN: 对每组 (C/G, H, W) 归一化 → 完全独立于 batch
y_gn = (x - μ_group) / √(σ²_group + ε) × γ + β   # μ,σ² 在每个样本内计算''')

doc.add_heading('GN 如何系统性解决上述 6 个问题', level=3)

add_table(doc,
    ['失效模式', 'BN 的行为', 'GN 的行为', '改善程度'],
    [
        ['模式1: SNR混合', 'running stats混合所有SNR\n测试时单个SNR不匹配', '每个样本独立归一化\n自适应于当前SNR的特征分布', '根本性解决'],
        ['模式2: 训练/推理不一致', '训练用batch stats\n推理用running stats\n两者来自不同分布', '训练和推理行为完全一致\n都是逐样本归一化', '根本性解决'],
        ['模式3: batch_size敏感性', 'batch_size=24时\n统计估计有噪声', '完全不依赖batch_size\n单样本即可稳定工作', '根本性解决'],
        ['模式4: 跨分辨率泛化', 'running stats基于256×256\n768×512可能分布不同', '逐样本归一化\n自动适应任何分辨率', '根本性解决'],
        ['模式5: 编解码器不对称', '解码器BN必须适应\n复杂噪声混合分布', '解码器逐样本自适应\n不需维护全局统计', '显著改善'],
        ['模式6: PReLU交互', '归一化偏差导致\nPReLU工作点偏移', '一致的归一化\nPReLU行为稳定', '显著改善'],
    ],
    col_widths=[2.5, 3.5, 3.5, 2]
)

doc.add_heading('GN 的潜在劣势与应对', level=3)
doc.add_paragraph(
    'GN 的主要弱点是：统计量仅在单样本的通道组内计算，对于特征图较小的情况，'
    '统计估计的样本量有限。以本项目 Layer 3 (16×16×1024, G=32 组) 为例：'
)
doc.add_paragraph('• BN：24×16×16 = 6144 个值 → 统计估计充分')
doc.add_paragraph('• GN：每组 (1024/32)×16×16 = 32×256 = 8192 个值 → 统计估计同样充分！')
doc.add_paragraph(
    '实际上，由于 Layer 3 的特征通道数很大（1024），即使分组后每组仍有 32 个通道，'
    '乘以 16×16 的空间尺寸，每组的样本量（8192）甚至超过 BN 的总样本量。GN 在这个场景下不会比 BN 差。'
)
doc.add_paragraph(
    '推荐 GN 配置：groups = 32（或 min(32, C) 对于通道数较小的层），'
    '这样每组的通道数约为 C/32，结合空间维度提供充足的统计样本。'
)

doc.add_page_break()

# ================================================================
# 第三部分
# ================================================================
doc.add_heading('第三部分：SimVQ 模块深度解析', level=1)

doc.add_heading('3.1 SimVQ 与传统 VQ-VAE 的架构对比', level=2)

add_table(doc,
    ['设计维度', '传统 VQ-VAE', 'SimVQ (本项目)', '差异评估'],
    [
        ['码本表示', 'E ∈ R^{K×D}，全可训练', 'E₀ ∈ R^{K×D} 冻结 + P ∈ R^{D×D} 可训练', 'SimVQ多一层间接性'],
        ['码本初始化', '随机 / K-Means', 'N(0, 1/√D) 随机 + 固定', '传统方法更优(可选K-Means)'],
        ['码本更新机制', 'EMA：e_new = β·e_old+(1-β)·z_enc', '全局投影：E = E₀P^T，P通过梯度更新', '传统方法更灵活'],
        ['码本坍缩防护', 'EMA + 码字重置', '冻结底层嵌入防止坍缩', 'SimVQ有优势(结构上防止)'],
        ['表达自由度/码字', 'D 个自由度(独立参数)', '等效 D 个自由度(但受全局约束)', '理论等价，实际受限'],
        ['训练稳定性', '需要调 EMA decay + 码字重置', '只需调 learning rate', 'SimVQ更简单'],
    ],
    col_widths=[2.5, 3.5, 3.5, 3]
)

doc.add_heading('3.2 冻结嵌入 + 投影层的数学局限性', level=2)

doc.add_paragraph('以下从线性代数和信息几何角度分析 SimVQ 设计的表达能力上限。')

doc.add_heading('引理 1：投影层的等效变换', level=3)
doc.add_paragraph(
    '设冻结嵌入矩阵 E₀ ∈ R^{K×D}（K=64, D=1024），投影矩阵 P ∈ R^{D×D}。'
    '有效码本为 E = E₀P^T ∈ R^{K×D}。E 的第 k 个码字为 e_k = P × (E₀的第k行)^T。'
)
doc.add_paragraph(
    '由于 K << D（64 << 1024），E₀ 的行向量最多张成一个 64 维子空间 S ⊂ R^D。'
    '投影矩阵 P 在 S 内的作用决定了有效码本的几何结构，P 在 S 的正交补 S^⊥ 内的作用是冗余的。'
)
doc.add_paragraph(
    '这意味着 P 的 D² = 1,048,576 个参数中，实际上只有 K×D = 65,536 个参数对码本几何有影响。'
    '其余 ~98 万个参数在训练中接收到的有效梯度极小（因为它们作用于随机初始化的、数据无关的方向）。'
)
doc.add_paragraph(
    '这解释了为什么 CODEBOOK_PROJ_LR = 1.75e-4 被设为普通学习率 (1.75e-5) 的 10 倍 —— '
    '代码作者可能已经发现 P 需要更大的学习率才能有效训练，但根本原因是 P 的有效自由度远小于名义参数数量。'
)

doc.add_heading('引理 2：表达能力的几何约束', level=3)
doc.add_paragraph(
    '传统 VQ 中，每个码字 e_k 可以独立地移动到其分配的编码器输出特征的均值位置：'
)
add_code(doc, '''e_k^{new} = (1-β) × e_k^{old} + β × mean({z_enc : assigned(z_enc)=k})''')
doc.add_paragraph(
    '这意味着码本可以适应任意的聚类结构——只要编码器输出的特征形成了 K 个簇，码字就能移动到这些簇的中心。'
)
doc.add_paragraph(
    '而 SimVQ 中，所有码字通过同一个 P 变换。如果最优码字位置为 {e_k*}，则 SimVQ 需要满足：'
)
add_code(doc, '''∃ P: E₀ P^T ≈ E*   即   e_k* ≈ P × (e₀)_k   ∀k''')
doc.add_paragraph(
    '这意味着最优码字必须能通过对初始随机向量集施加同一个线性变换得到。'
    '但编码器特征 Z_enc 在训练过程中形成的聚类结构，与随机初始化的 E₀ 行向量施加全局线性变换后的结果之间，'
    '不能保证存在一个足够好的 P 使它近似。'
)

doc.add_heading('量化分析：重构误差下界', level=3)
doc.add_paragraph('定义：')
add_code(doc, '''传统VQ: L_vq* = min_E E_z[ min_k ||z - e_k||² ]
SimVQ:  L_svq* = min_P E_z[ min_k ||z - (E₀P^T)_k||² ]''')
doc.add_paragraph(
    '显然 L_svq* ≥ L_vq*，因为 SimVQ 的可行域是传统 VQ 可行域的子集。'
    '差距 Δ = L_svq* - L_vq* 取决于最优码本 E* 能否被 E₀P^T 很好地近似。'
)
doc.add_paragraph(
    '当 E₀ 随机初始化时，其行向量近似正交且均匀分布在球面上。P 可以将这个正交基旋转和缩放，'
    '但不能改变行向量间的"正交"拓扑关系。因此，如果数据的最优聚类中心不具有近似正交的结构，'
    'SimVQ 就存在不可消除的表达误差。这个误差成为系统性能的一个下界。'
)

doc.add_heading('3.3 K << D 高维稀疏码本的几何分析', level=2)

doc.add_paragraph(
    '对于 Layer 3（K=64, D=1024），我们分析其几何特性：'
)

doc.add_heading('随机初始化后的码本几何', level=3)
doc.add_paragraph(
    'E₀ 的行向量 e₀_k ~ N(0, (1/D)I)，即每个元素是独立高斯 N(0, 1/D)。这个初始化有以下性质：'
)
doc.add_paragraph('• 期望范数：E[||e₀_k||²] = D × (1/D) = 1（单位长度，近似）')
doc.add_paragraph(
    '• 期望内积：E[⟨e₀_i, e₀_j⟩] = 0 for i≠j → 向量近似正交'
)
doc.add_paragraph(
    '• 余弦相似度：E[cos(e₀_i, e₀_j)] ≈ 1/√D ≈ 0.031 → 几乎不相关'
)
doc.add_paragraph(
    '• 码字间最小距离：min_{i≠j} ||e₀_i - e₀_j||² ≈ 2 - 2/√D ≈ 2 → 距离较大'
)

doc.add_heading('投影后的码本变化', level=3)
doc.add_paragraph(
    'P^T 作用相当于对 E₀ 的每一行做相同的线性变换。设 P 的 SVD 分解为 P = U Σ V^T。'
    '则 P^T = V Σ U^T。对 E₀ 的作用是：先旋转（U^T），再缩放（Σ），再旋转（V）。'
)
doc.add_paragraph(
    '旋转不改变向量间的角度和距离。缩放（Σ 的对角线元素）会改变码字在某些方向上的长度，'
    '从而改变码字间的距离矩阵。总的来说，P 能做的变换有：'
)
doc.add_paragraph('• 全局旋转：改变码本的整体朝向 → 帮助对齐数据的主要方向')
doc.add_paragraph('• 各向异性缩放：在数据变化大的方向拉长，变化小的方向缩短 → 适应数据的协方差结构')
doc.add_paragraph('• 不能做：改变码字之间的相对拓扑关系、使两个原本远离的码字变得很近、改变码字的数量')
doc.add_paragraph(
    '结论：SimVQ 的码本只能实现一个"全局仿射变换"程度的适应。对于复杂的自然图像数据（Kodak），'
    '仅靠全局线性变换极难让 64 个随机初始化的码字有效覆盖所有可能的 16×16 局部特征模式。'
)

doc.add_heading('3.4 码本坍缩风险的深层机制', level=2)

doc.add_paragraph(
    '传统 VQ 的主要坍缩机制：某些码字接收不到足够的编码器输出分配，EMA 更新缓慢，'
    '最终这些"死码字"被"活码字"拉近（因为 EMA 的指数衰减特性），导致多个码字坍缩到同一位置。'
)
doc.add_paragraph(
    'SimVQ 的冻结嵌入设计从结构上防止了这种坍缩——即使某个码字从未被使用，'
    '它的底层嵌入向量也不会改变。但 SimVQ 引入了另一种坍缩风险：投影坍缩。'
)

doc.add_heading('投影坍缩（Projection Collapse）', level=3)
doc.add_paragraph(
    '由于所有码字共享同一个投影矩阵 P，如果 P 的某些奇异值变得非常小，'
    '所有码字在对应方向上的分量都会被压缩。极端情况下，如果 P 的秩下降（某些奇异值→0），'
    '所有码字被投影到一个低维子空间，等效于码本容量大幅缩小。'
)
doc.add_paragraph(
    '监控方法：定期计算 P 的有效秩（奇异值谱），如果有效秩持续下降，说明发生了投影坍缩。'
)
add_code(doc, '''def compute_effective_rank(proj_weight):
    s = torch.linalg.svdvals(proj_weight)
    s_normalized = s / s.sum()
    entropy = -(s_normalized * torch.log(s_normalized)).sum()
    return torch.exp(entropy).item()  # 有效秩''')

doc.add_heading('使用率坍缩（Utilization Collapse）', level=3)
doc.add_paragraph(
    '即使码字不坍缩（几何上保持分离），也可能出现使用率坍缩——大部分编码器输出只分配到少数码字。'
    '当 K=64 但只有 10 个码字被使用时，等效码本大小就是 10。在测试链路的 LDPC 错误面前，'
    '10 个码字之间的汉明距离可能不足以提供错误保护。'
)
doc.add_paragraph(
    'cfg.COMMITMENT_COST = 0.25 控制编码器对码字的"承诺"强度。值越小，编码器越"不愿意"改变自己的输出'
    '来靠近码字，容易导致使用率坍缩。建议增大到 0.5-1.0 以鼓励更均匀的码本使用。'
)

doc.add_page_break()

# ================================================================
# 第四部分
# ================================================================
doc.add_heading('第四部分：逐模块微观问题清单', level=1)

doc.add_heading('4.1 编码器设计问题', level=2)

doc.add_heading('问题 E1：ResidualBlock 中缺少最后的激活函数', level=3)
add_code(doc, '''class ResidualBlock(nn.Module):
    def forward(self, x):
        identity = x
        out = self.conv1(x)
        out = self.bn(out)
        out = self.prelu(out)
        out = self.conv2(out)       # conv2后没有BN，也没有激活
        out = out + identity         # 直接加残差
        return out                   # ← 输出没有经过非线性激活!''')
doc.add_paragraph(
    '标准的 Pre-activation ResBlock 会在 conv2 和加法后接 BN+ReLU。当前设计在 conv2 后既没有 BN 也没有激活。'
    '虽然这不是错误（原始的 ResNet 论文中 bottleneck block 的最后一个 conv 后确实没有激活），但：'
)
doc.add_paragraph('• conv2 的输出范围不受控制（无 BN 归一化），可能导致残差分支的数值量级与 identity 分支不匹配')
doc.add_paragraph('• 缺少最终的激活使得残差块的输出可以取任意实数值，与后续模块的期望输入分布可能不一致')

doc.add_heading('问题 E2：DownSampleBlock 中多余的独立 BN+PReLU', level=3)
add_code(doc, '''class DownSampleBlock(nn.Module):
    def forward(self, x):
        x = self.res1(x)       # 包含 conv1→BN→PReLU→conv2
        x = self.down(x)       # stride=2 下采样卷积
        x = self.res2(x)       # 包含 conv1→BN→PReLU→conv2
        x = self.bn(x)         # ← 独立的BN (与res2内部的BN功能重复)
        x = self.act(x)        # ← 独立的PReLU
        x = self.tail(x)       # 额外的卷积
        return x''')
doc.add_paragraph(
    'ResidualBlock.res2 的 conv2 之后没有 BN，但 DownSampleBlock 在 res2 之后又加了一个 BN+PReLU。'
    '这种"双 BN + 双 PReLU"的设计（res2 内部 1 个 BN + 外部 1 个 BN，res2 内部 1 个 PReLU + 外部 1 个 PReLU）'
    '增加了参数量和计算量，但收益不明确。更标准的设计是在 res2 的 conv2 后加 BN，去掉外部的独立 BN。'
)

doc.add_heading('4.2 解码器设计问题', level=2)

doc.add_heading('问题 D1：Nearest-Neighbor 上采样产生块效应', level=3)
add_code(doc, '''x = F.interpolate(x, scale_factor=2, mode='nearest')''')
doc.add_paragraph(
    'Nearest-neighbor 上采样通过复制像素值实现 2× 放大。这会在特征图中引入"块状"伪影——'
    '相邻 2×2 区域内的值完全相同。后续的 3×3 卷积可以部分平滑这些块效应，但无法完全消除。'
)
doc.add_paragraph(
    '对于经过 VQ 量化 + 信道损坏的特征，块效应可能被放大：如果两个相邻位置被量化到不同的码字，'
    '上采样后形成的 2×2 块就会有截然不同的值，在边界处产生高频伪影。'
)
doc.add_paragraph(
    '建议：将 up_mode 改为 "bilinear"（双线性插值），它产生平滑的上采样结果，'
    '且对梯度流更友好（双线性插值是逐段线性的，梯度更平滑）。'
    '注意需要设置 align_corners=False。'
)

doc.add_heading('问题 D2：Skip Connection Dropout 的训练/推理 gap', level=3)
doc.add_paragraph(
    'SKIP_DROPOUT_P = [0.5, 0.45, 0.05]。训练时 Layer 0 的 skip 有 50% 概率被完全丢弃。'
    '这强制解码器学习不依赖浅层特征也能重建，但也导致了一个 gap：'
)
doc.add_paragraph(
    '• 训练时：解码器在"有 skip"和"无 skip"两种模式间切换，学到的权重是两种模式的折中'
)
doc.add_paragraph(
    '• 推理时：skip 总是 100% 通过，解码器突然获得大量它不习惯使用的信息'
)
doc.add_paragraph(
    '这种 gap 的典型表现是：推理时重建图像出现"过度锐化"或"纹理闪烁"——'
    '因为 skip 提供了训练时一半时间不存在的细节，解码器不知如何处理。'
    '建议在推理时也使用一定概率的 skip dropout（如 p=0.1），或者改用'
    '更温和的正则化方式（如在 skip 特征上添加高斯噪声而非完全丢弃）。'
)

doc.add_heading('问题 D3：最深层的瓶颈设计', level=3)
doc.add_paragraph(
    '解码器的输入从最深层开始（16×16×1024 或 48×32×1024）。这个最深层特征经过 VQ 量化后，'
    '包含了图像的最高层语义信息，但空间分辨率最低。整个解码器需要从这 256 个空间位置'
    '（16×16）重建出 256×256 的高分辨率图像。这意味着每个空间位置承载了 64×64 像素区域的信息——'
    '信息压缩比极高。如果最深层 VQ 有任何错误（码字选错），64×64 的整个区域都会受影响。'
)
doc.add_paragraph(
    '这是 VQ 信源编码的固有问题，但在这个架构中尤为突出，因为最深层:'
    '(a) 码本容量最小（64 码字覆盖最高语义），(b) VQ 损失权重最大（×10），'
    '(c) 信道错误的影响也最大（1 个索引错误影响 4096 像素）。'
)

doc.add_heading('4.3 信道模块设计问题', level=2)

doc.add_heading('问题 C1：训练时的信道编码速率与实际 LDPC 不匹配', level=3)
doc.add_paragraph(
    'CHANNEL_CODING_RATE_TRAIN = 0.5，LDPC_R = 0.5。速率数值一致，但实际含义不同：'
)
doc.add_paragraph('• 训练中：速率 0.5 意味着"每 2 个传输比特中 1 个是信息比特"，这直接被用在 BER 公式中')
doc.add_paragraph(
    '• 测试中：速率 0.5 意味着 128 个信息比特被编码为 256 个编码比特。LDPC 的纠错能力不仅取决于速率，'
    '还取决于具体的校验矩阵结构和 BP 译码算法'
)
doc.add_paragraph(
    '同等速率下，LDPC 的实际纠错性能可能与理论 BER 公式相差 1-2 dB（在 10^-3 BER 水平）。'
    '这 1-2 dB 的差异直接转化为训练和测试之间的系统性偏差。'
)

doc.add_heading('问题 C2：测试时 LDPC 译码失败无检测机制', level=3)
doc.add_paragraph(
    'test_real.py 中的 LDPC 译码器（基于 Sionna 的 5G LDPC 译码器）通过 BP 算法迭代译码。'
    '在低 SNR 下，BP 算法可能不收敛，输出的"译码比特"包含大量错误。但代码中没有检查：'
)
add_code(doc, '''# 当前：直接使用译码结果，不验证
decoded_bits = ldpc_decode(llrs.cpu().numpy(), ldpc_code)

# 应该：检查译码是否成功（通过校验矩阵验证）
# 5G LDPC编码是系统码，可以通过H×c^T=0验证
# 如果译码失败，至少知道哪些bits不可靠''')
doc.add_paragraph(
    '没有译码成功性验证意味着：一些完全错误的译码结果被当作"正确"结果送入解码器，'
    '产生极差的重建图像，拉低了平均 MS-SSIM/PSNR。这些实际上是"译码失败"而不是"重建失败"。'
)

doc.add_heading('4.4 损失函数设计问题', level=2)

doc.add_heading('问题 L1：MSE 对 MS-SSIM 的间接优化', level=3)
doc.add_paragraph(
    '这是一个根本性的指标-损失不匹配问题。MSE 最小化等价于最大化 PSNR。MS-SSIM 则衡量结构相似性。'
    '虽然两者有一定正相关性，但在以下情况下会分离：'
)
doc.add_paragraph('• 模糊平滑的图像：MSE 较低（因为像素值接近均值），但 MS-SSIM 很低（缺乏结构）')
doc.add_paragraph('• 有噪声但结构正确的图像：MSE 较高，但 MS-SSIM 可以很高')
doc.add_paragraph(
    '在这个项目中，经过信道传输后的重建图像往往偏模糊（MSE 训练的固有倾向），'
    '这直接导致 MS-SSIM 指标不理想。'
)

doc.add_heading('问题 L2：VQ 损失权重的标量问题', level=3)
doc.add_paragraph(
    'LAYER_LOSS_WEIGHTS = [1, 1, 5, 10]。深层权重 10 倍于浅层。但不同层的 VQ 损失天然具有不同的量级：'
)
doc.add_paragraph('• Layer 0：维度 128，特征图 128×128 → VQ 损失在 16384 个位置 × 128 维上求和')
doc.add_paragraph('• Layer 3：维度 1024，特征图 16×16 → VQ 损失在 256 个位置 × 1024 维上求和')
doc.add_paragraph(
    '实际上，由于 embedding_dim 差异（1024 vs 128 = 8×），Layer 3 每个空间位置的 VQ 损失本身就是 Layer 0 的 ~8 倍。'
    '再乘以 10 的权重，Layer 3 的 VQ 损失贡献是 Layer 0 的 ~80 倍。这几乎完全压制了浅层的 VQ 训练信号。'
)
doc.add_paragraph(
    '建议：在分配权重前，先对各层 VQ 损失按其自然量级做归一化（如除以 embedding_dim），再应用语义权重。'
)

doc.add_heading('4.5 数据处理问题', level=2)

doc.add_heading('问题 D1：训练/测试数据集域偏移', level=3)
doc.add_paragraph(
    'Cars196 包含 196 类汽车，背景相对简单（路面、展厅等）。Kodak 是通用图像集，包含人像、风景、建筑等。'
    '这种域偏移对 VQ 码本影响尤其大——码本学到的"典型特征"是汽车的纹理（金属漆面、轮胎纹路、车窗反光），'
    '但测试时需要量化人脸皮肤、树叶、天空等完全不同纹理的特征。'
)
doc.add_paragraph(
    '由于 SimVQ 的码本表达能力本就受限（见第三部分），跨域泛化的难度进一步加大。'
)

doc.add_heading('问题 D2：测试图像 Resize 导致的信息损失', level=3)
add_code(doc, '''elif mode == 'test':
    transform = transforms.Compose([
        transforms.Resize((768, 512)),  # 强制缩放，改变宽高比
        ...''')
doc.add_paragraph(
    'Kodak 数据集的图像原始尺寸为 768×512（或 512×768），但 test 模式下使用了 Resize((768, 512))，'
    '这会强制将所有图像变为横向 768×512 的尺寸。对于原始竖向的图像（512×768），这会导致严重的宽高比失真，'
    '将影响 MS-SSIM 的计算（因为参考图像也被同样形变）。'
)

doc.add_page_break()

# ================================================================
# 第五部分
# ================================================================
doc.add_heading('第五部分：快速验证方法论 —— 如何在 10-50 Epoch 内判断改动有效性', level=1)

doc.add_paragraph(
    '目前训练需要 400 epochs，完整训练成本高。以下设计一套代理指标体系，'
    '使改动效果能在 10-50 epochs 内得到可靠判断，避免"训练完才发现没用"的问题。'
)

doc.add_heading('5.1 代理指标体系（6 个核心指标）', level=2)

doc.add_heading('指标 1：真实链路快速验证 MS-SSIM (Gold Standard)', level=3)
doc.add_paragraph(
    '这是最直接的代理指标。每 5 个 epoch，用 20 张验证图像跑完整的 test_real.py 链路（BPSK+LDPC+AWGN），'
    '在 2 个代表性 SNR（6dB 和 12dB）测量 MS-SSIM。'
)
doc.add_paragraph('• 20 张图 × 2 SNR × LDPC ≈ 30 秒 → 每 5 epoch 评估一次，开销可接受')
doc.add_paragraph('• 6dB 代表"困难"条件，12dB 代表"良好"条件 → 覆盖测试范围')
doc.add_paragraph('• 如果 10 epoch 内这个指标持续上升 → 改动有效')
doc.add_paragraph('• 如果 10 epoch 内几乎不变 → 改动无效或学习率太低')

doc.add_heading('指标 2：码本利用率收敛速度', level=3)
doc.add_paragraph(
    '每 5 个 epoch 统计一次各层 active_ratio。关键看收敛速度而非绝对值：'
)
doc.add_paragraph('• 理想：前 10 epoch active_ratio 迅速达到 >80% → 码本训练健康')
doc.add_paragraph('• 警告：10 epoch 后 active_ratio <50% → VQ 存在坍缩，码本容量浪费')
doc.add_paragraph('• 危险：active_ratio 持续下降 → 投影坍缩（见 3.4 节）')
doc.add_paragraph(
    '特别关注 Layer 2 和 Layer 3（权重最高的两层），如果这两层的 active_ratio 低，'
    '说明最关键的语义层码本未能充分利用。'
)

doc.add_heading('指标 3：量化失真率 (Pre-Channel Quantization Error)', level=3)
doc.add_paragraph(
    '在加信道噪声之前，测量编码器特征与 VQ 量化特征的 MSE。这隔离了 VQ 本身的性能：'
)
add_code(doc, '''with torch.no_grad():
    for feat in encoder_features:
        _, quantized, _ = vq(feat)
        q_error += F.mse_loss(quantized, feat).item()''')
doc.add_paragraph('• 如果这个值在 10 epoch 后稳定且较低 → VQ 质量好，瓶颈在信道')
doc.add_paragraph('• 如果这个值居高不下 → VQ 容量不足（K 太小）或 SimVQ 投影表达能力不够')
doc.add_paragraph('• 改动后这个值明显降低 → 改动改善了 VQ 质量（如增大码本、改进初始化）')

doc.add_heading('指标 4：信道鲁棒性差距 (Channel Robustness Gap)', level=3)
doc.add_paragraph(
    '分别用 SNR=20dB（近似无噪）和 SNR=6dB 做重建，计算 MS-SSIM 的差值：'
)
add_code(doc, '''robustness_gap = MS-SSIM(SNR=20dB) - MS-SSIM(SNR=6dB)''')
doc.add_paragraph('• gap 越小 → 模型越鲁棒 → 训练的信道模拟越接近真实')
doc.add_paragraph('• 如果 gap 在训练过程中持续缩小 → 模型正在学习信道鲁棒性（好！）')
doc.add_paragraph('• 如果 gap 不变 → 模型没有学到信道鲁棒性（STE 问题的直接证据）')
doc.add_paragraph(
    '这个指标直接反映了"训练信道模拟质量"——如果改动了信道模块，应该看到 gap 明显缩小。'
)

doc.add_heading('指标 5：梯度 SNR (Gradient Signal-to-Noise Ratio)', level=3)
doc.add_paragraph(
    '每 100 步，计算编码器梯度和解码器梯度的范数比：'
)
add_code(doc, '''encoder_grad_norm = sum(p.grad.norm() for p in encoder.parameters())
decoder_grad_norm = sum(p.grad.norm() for p in decoder.parameters())
grad_ratio = encoder_grad_norm / (decoder_grad_norm + 1e-8)''')
doc.add_paragraph('• grad_ratio ≈ 0.5-2.0 → 编解码器学习速度匹配 → 好')
doc.add_paragraph('• grad_ratio << 0.1 → 编码器几乎不学习 → STE 截断了编码器梯度（问题2的直接证据）')
doc.add_paragraph('• grad_ratio 随时间变化 → 如果趋近于 0，编码器在"放弃"学习')
doc.add_paragraph(
    '这个指标对于评估"是否更换了 STE"或"是否使用了可微分信道"特别有用——'
    '如果改动有效，grad_ratio 应该明显上升并稳定。'
)

doc.add_heading('指标 6：特征分布偏移 (Feature Distribution Shift)', level=3)
doc.add_paragraph(
    '测量解码器各层输入特征的逐通道均值/方差，对比训练集和少量测试集图像：'
)
add_code(doc, '''# 训练图像 (经过训练信道)
feat_train = decoder_layer_output(train_img, training_channel)
# 测试图像 (经过测试信道)
feat_test = decoder_layer_output(test_img, test_channel)
# 计算分布距离
shift = wasserstein_distance(feat_train, feat_test)''')
doc.add_paragraph('• shift 小 → 训练信道模拟准确 → 好')
doc.add_paragraph('• shift 大 → 训练和测试的特征分布系统性不同 → BN 失效 + 解码器泛化差')
doc.add_paragraph('• 如果改为 GN 后 shift 应该变小（因为不再依赖 running statistics）')

doc.add_heading('5.2 实用评估协议', level=2)

doc.add_paragraph('以下是一套可操作的评估流程，适合在改动后 10-50 epoch 内判断方向：')

doc.add_heading('Phase 1：基线记录（改动前，耗时 ~5 分钟）', level=3)
doc.add_paragraph('1. 在 checkpoint 上运行 20 张验证图的真实链路评估，记录 SNR=6dB 和 12dB 的 MS-SSIM')
doc.add_paragraph('2. 计算并记录所有 6 个代理指标的基线值')
doc.add_paragraph('3. 保存一份可视化：2-3 张验证图的"原图 vs 重建图"对比')

doc.add_heading('Phase 2：快速训练 + 密集监控（改动后 0-20 epochs，耗时 ~2 小时）', level=3)
doc.add_paragraph('1. 使用 2-3× 正常学习率（1e-4 级别），训练 20 个 epoch')
doc.add_paragraph('2. 每 5 个 epoch 评估指标 1（真实链路 MS-SSIM）+ 指标 2（码本利用率）')
doc.add_paragraph('3. 每 100 step 记录指标 5（梯度 SNR）')
doc.add_paragraph('4. 在第 10 和 20 epoch 评估全部 6 个指标')

doc.add_heading('Phase 3：决策判断（epoch 20）', level=3)
add_table(doc,
    ['判断条件', '结论', '后续动作'],
    [
        ['指标1 ↑ > 0.02 + 指标2 > 70%', '改动非常有效', '继续训练至50 epoch确认，降低LR微调'],
        ['指标1 ↑ 0.01-0.02 + 指标2 > 50%', '改动有正向效果', '继续训练观察，可能需要调整超参'],
        ['指标1 几乎不变 (±0.005)', '改动无效或LR太低', '尝试调整LR，如果仍无效则放弃'],
        ['指标1 ↓', '改动有害', '回滚，分析原因'],
        ['指标4 (robustness gap) ↓ 明显', '信道模拟改善显著', '这是最重要的积极信号'],
        ['指标5 (grad_ratio) 上升', '编码器开始学习鲁棒性', 'STE修复有效！'],
    ],
    col_widths=[3.5, 2.5, 5]
)

doc.add_heading('Phase 4：确认阶段（epoch 20-50，耗时 ~4 小时）', level=3)
doc.add_paragraph(
    '如果 Phase 3 判断有效，降低学习率到正常水平（1.75e-5），继续训练到 50 epoch，'
    '在第 30、40、50 epoch 各做一次完整的 test_real.py 评估（所有 5 个 SNR），确认趋势。'
)

doc.add_heading('5.3 各指标的早期预警阈值', level=2)

add_table(doc,
    ['指标', '健康范围 (epoch 10)', '警告阈值', '危险阈值', '测量频率'],
    [
        ['真实链路 MS-SSIM(12dB)', '>0.80', '<0.75', '<0.70', '每5 epoch'],
        ['码本 active_ratio (平均)', '>75%', '<60%', '<40%', '每5 epoch'],
        ['量化失真 MSE (归一化)', '<0.05', '0.05-0.10', '>0.10', '每5 epoch'],
        ['鲁棒性差距 gap', '<0.15', '0.15-0.25', '>0.25', '每5 epoch'],
        ['梯度比 grad_ratio', '0.3-2.0', '<0.1', '<0.05', '每100 step'],
        ['特征偏移 shift', '<0.5σ', '0.5-1.0σ', '>1.0σ', '每10 epoch'],
    ],
    col_widths=[3.2, 2.5, 2, 2, 2.5]
)

doc.add_page_break()

# ================================================================
# 附录：关键代码位置速查
# ================================================================
doc.add_heading('附录：关键问题与代码位置速查表', level=1)

add_table(doc,
    ['问题编号', '问题简述', '关键文件:行号', '建议改动方向'],
    [
        ['P1', '训练/测试信道不一致', 'models/channel.py:41-65\nmodels/deepsc.py:69-80\ntest_real.py:137-153', '替换为可微分AWGN信道\n训练中包含BPSK模拟'],
        ['P2', 'STE梯度截断', 'models/deepsc.py:79', '用软量化+Gumbel-Softmax\n替换.detach()模式'],
        ['P3', '调制方式不匹配', 'models/deepsc.py:46-52\ntest_real.py:141', '固定训练mod_bits=1(BPSK)'],
        ['P4', 'MSE损失与MS-SSIM\n指标不一致', 'losses/deepsc_loss.py:16', '增加LPIPS感知损失\n或MS-SSIM损失'],
        ['P5', '码本容量不足', 'config.py:12-13', '增大K至128-256\n增加commitment_cost至0.5-1.0'],
        ['P6', '训练/测试数据域偏移', 'data/datasets.py:42-62\nconfig.py:44-46', '增加训练数据多样性\n或多分辨率训练'],
        ['P7', 'VQ损失权重失衡', 'config.py:17', '按embedding_dim归一化\n后重新分配权重'],
        ['P8', 'BN→GN替换', 'models/semantic_encoder.py\nmodels/semantic_decoder.py', '20个BN层→GN(G=32)'],
        ['P9', 'Nearest上采样块效应', 'models/semantic_decoder.py:53', '改为bilinear上采样'],
        ['P10', 'Skip Dropout\n训练/推理gap', 'config.py:21\nmodels/semantic_decoder.py:16', '推理时保留小概率dropout\n或改用高斯噪声'],
        ['P11', 'LDPC译码失败无检测', 'test_real.py:149', '增加CRC校验或\nH矩阵验证'],
        ['P12', 'BN running statistics\nSNR混合', '所有BN层\n(共20处)', '替换为GN根本性解决'],
        ['P13', 'SimVQ投影层\n有效秩下降', 'models/vector_quantizer.py\n:10-25', '监控proj奇异值谱\n必要时正则化'],
        ['P14', 'Clamping引入\n码字偏差', 'models/channel.py:63', '用模运算替代clamp\n或重新采样'],
    ],
    col_widths=[1, 2.5, 3.5, 3.5]
)

# ===== 保存 =====
doc.save(OUTPUT_PATH)
print(f"深度分析报告已生成: {OUTPUT_PATH}")
print(f"文件大小: {os.path.getsize(OUTPUT_PATH) / 1024:.1f} KB")
