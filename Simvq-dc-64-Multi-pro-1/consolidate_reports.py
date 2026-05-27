"""
为三份 SimVQ 报告添加交叉引用，使其形成有机整体。
使用 python-docx 直接编辑 docx 文件。
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import copy
from datetime import datetime

BASE = r"D:\Project\Pycharm_project_new\Simvq-dc-64-Multi-pro"


def add_paragraph_after(doc, search_text, new_text, bold=False, italic=False, style=None):
    """在包含 search_text 的段落后插入新段落。返回插入的段落或 None。"""
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph
    for i, para in enumerate(doc.paragraphs):
        if search_text in para.text:
            # Create a new paragraph element and insert it after the current one
            new_p = OxmlElement('w:p')
            para._element.addnext(new_p)
            new_para = Paragraph(new_p, para._parent)
            # Set text
            if new_text:
                new_para.add_run(new_text)
            if style:
                new_para.style = style
            if bold or italic:
                for run in new_para.runs:
                    run.bold = bold
                    run.italic = italic
            return new_para
    return None


def add_paragraph_before(doc, search_text, new_text, bold=False, italic=False, style=None):
    """在包含 search_text 的段落前插入新段落。"""
    for i, para in enumerate(doc.paragraphs):
        if search_text in para.text:
            # Insert before: add to the previous paragraph's XML
            new_para = para.insert_paragraph_before(new_text)
            if style:
                new_para.style = style
            if bold or italic:
                for run in new_para.runs:
                    run.bold = bold
                    run.italic = italic
            return new_para
    return None


def add_related_reports_section(doc, report_type):
    """在报告末尾添加'相关报告'章节。"""
    # Find the last meaningful paragraph
    last_idx = len(doc.paragraphs) - 1
    for i in range(len(doc.paragraphs) - 1, -1, -1):
        if doc.paragraphs[i].text.strip():
            last_idx = i
            break

    if report_type == "v1":
        content = [
            ("相关报告", True),
            ("本报告是 SimVQ 语义通信系统性能分析系列报告的第一部分（总览篇），建议配合以下报告阅读：", False),
            ("", False),
            ("SimVQ_深度性能分析报告_v2.docx —— 深度分析篇", True),
            ("涵盖：训练/测试信道鸿沟的微观解剖、BN vs GN 选型分析（20层BN失效模式）、SimVQ 数学局限性（冻结嵌入+投影层）、逐模块微观问题审查（编码器/解码器/信道/损失函数/数据）、快速验证方法论（10-50 Epoch 判断改动有效性）。建议在阅读本报告的问题诊断后，参考 v2 获取深层技术分析。", False),
            ("", False),
            ("SimVQ_遗漏问题补充报告_v3.docx —— 补充篇", True),
            ("涵盖：11 个前两份报告未覆盖的新问题（含验证管线缺陷 N1、解码器输出无界 N2、Adam β₁ 训练不稳定 N3、编码器/码本尺度不匹配 N4 等），以及三版报告的 33 问题总索引。本报告提出的方案 B 未考虑 v3 中的新发现，实际实施时建议综合 v3 的修正方向。", False),
        ]
    elif report_type == "v2":
        content = [
            ("相关报告", True),
            ("本报告是 SimVQ 语义通信系统性能分析系列报告的第二部分（深度分析篇），建议配合以下报告阅读：", False),
            ("", False),
            ("SimVQ性能分析报告.docx —— 总览篇", True),
            ("涵盖：项目概述、8 个核心问题诊断（P1-P8）、3 套优化方案对比（A/B/C）、方案 B 详细实施说明。建议先阅读该报告建立全局认知，再阅读本报告获取深层技术分析。", False),
            ("", False),
            ("SimVQ_遗漏问题补充报告_v3.docx —— 补充篇", True),
            ("涵盖：11 个前两份报告未覆盖的新问题（N1-N11），以及三版报告的 33 问题总索引。其中 N1（验证管线使用错误信道）的严重程度不亚于本报告分析的核心问题。", False),
        ]
    else:  # v3
        content = [
            ("相关报告", True),
            ("本报告是 SimVQ 语义通信系统性能分析系列报告的第三部分（补充篇），建议配合以下报告阅读：", False),
            ("", False),
            ("SimVQ性能分析报告.docx —— 总览篇", True),
            ("涵盖：项目概述、8 个核心问题诊断（P1-P8）、3 套优化方案对比（A/B/C）、方案 B 详细实施说明。本报告中的新问题为方案实施提供了额外的修正方向。", False),
            ("", False),
            ("SimVQ_深度性能分析报告_v2.docx —— 深度分析篇", True),
            ("涵盖：训练/测试信道鸿沟微观解剖、BN vs GN 选型分析（20层BN失效模式）、SimVQ 数学局限性、逐模块微观问题审查、快速验证方法论。本报告中的 N3（Adam β₁）、N4（尺度不匹配）、N8（投影层正则化）等问题与 v2 有深度技术关联。", False),
        ]

    # Add empty paragraph before section
    doc.add_paragraph("")
    # Add horizontal rule (using a paragraph with bottom border)
    hr = doc.add_paragraph("")
    pPr = hr._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '6',
        qn('w:space'): '1',
        qn('w:color'): '2E75B6'
    })
    pBdr.append(bottom)
    pPr.append(pBdr)

    for text, is_heading in content:
        p = doc.add_paragraph(text)
        if is_heading:
            p.style = doc.styles['Heading 1']
        else:
            for run in p.runs:
                run.font.size = Pt(11)


def process_v1():
    """处理 v1 报告"""
    path = f"{BASE}/SimVQ性能分析报告.docx"
    print(f"处理 {path}...")
    doc = Document(path)

    # 1. 在"一、项目概述"前插入"关于本报告"
    add_paragraph_before(
        doc,
        "一、项目概述",
        "关于本报告",
        bold=True
    )
    intro = add_paragraph_after(
        doc,
        "关于本报告",
        "本报告是 SimVQ 语义通信系统性能分析系列报告的第一部分（总览篇），完成于 2026年5月13日。报告识别了 8 个核心性能问题（P1-P8），提出了三套递进式优化方案（方案 A/B/C），并推荐方案 B 作为实施首选。本报告侧重问题诊断和方案设计，深层技术分析请参见系列第二份报告（SimVQ_深度性能分析报告_v2.docx），补充问题请参见第三份报告（SimVQ_遗漏问题补充报告_v3.docx）。",
        italic=True
    )
    if intro:
        for run in intro.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(100, 100, 100)

    # 2. 在"三、性能问题根因诊断"的问题8末尾添加补充说明
    #    找到"四、优化方案"前面
    found = False
    for i, para in enumerate(doc.paragraphs):
        if "四、优化方案（3套方案对比）" in para.text and not found:
            found = True
            # Insert before this paragraph
            note1 = para.insert_paragraph_before("")
            note2 = para.insert_paragraph_before(
                "补充说明：以上 8 个问题（P1-P8）为本报告的初始诊断结果。在后续深入分析中，系列第二份报告（v2）进一步揭示了 Batch Normalization 在此场景下的 6 个失效模式、SimVQ 冻结嵌入的数学局限性、以及逐模块微观设计问题。系列第三份报告（v3）额外发现了 11 个遗漏问题（N1-N11），其中 N1（验证管线使用错误信道导致最优模型从未被正确选出）的严重程度不亚于本报告的 CRITICAL 级别问题。建议综合三份报告的全部发现来制定优化策略。"
            )
            for run in note2.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(180, 100, 0)
                run.italic = True
            break

    # 3. 在方案对比表之后、五之前添加注释
    for i, para in enumerate(doc.paragraphs):
        if "五、推荐方案：方案 B 详细说明" in para.text:
            note = para.insert_paragraph_before("")
            note2 = para.insert_paragraph_before(
                "重要提示：方案 B 设计于 v1 报告阶段（2026年5月13日），当时尚未发现 v3 报告中的 11 个遗漏问题。实际实施时，强烈建议在方案 B 的 5 项改动基础上，额外考虑 v3 中以下关键修正：N1（验证管线改用真实链路评估选模）、N2（解码器末尾添加 Tanh 约束输出范围）、N3（Adam β₁ 恢复为 0.9）、N4（VQ 距离计算前做 L2 归一化）。这四项修正改动量小、风险低，但对最终效果影响显著。"
            )
            for run in note2.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(180, 80, 0)
                run.italic = True
            break

    # 4. 在末尾添加"相关报告"章节
    add_related_reports_section(doc, "v1")

    doc.save(path)
    print("v1 处理完成 ✓")


def process_v2():
    """处理 v2 报告"""
    path = f"{BASE}/SimVQ_深度性能分析报告_v2.docx"
    print(f"处理 {path}...")
    doc = Document(path)

    # 1. 在"第一部分"前插入"关于本报告"
    add_paragraph_before(
        doc,
        "第一部分：宏观架构审查",
        "关于本报告",
        bold=True
    )
    intro = add_paragraph_after(
        doc,
        "关于本报告",
        "本报告是 SimVQ 语义通信系统性能分析系列报告的第二部分（深度分析篇），完成于 2026年5月14日。报告在 v1（总览篇）的 8 个核心问题基础上进行了深度技术剖析，涵盖四大专题：训练/测试信道鸿沟的微观解剖、Batch Normalization vs Group Normalization 选型分析、SimVQ 模块数学局限性的形式化证明、以及逐模块微观设计审查。此外提供了快速验证方法论，支持在 10-50 epoch 内判断改动有效性。建议先阅读 v1（SimVQ性能分析报告.docx）建立全局认知，再阅读本报告。",
        italic=True
    )
    if intro:
        for run in intro.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(100, 100, 100)

    # 2. 在各部分末尾添加"相关内容"标注
    # 第一部分末尾 → 在"第二部分"前插入
    for i, para in enumerate(doc.paragraphs):
        if "第二部分：Normalization 层选型分析" in para.text:
            ref = para.insert_paragraph_before(
                "▸ 相关内容：本部分识别的问题（P1 训练/测试信道不一致、P2 STE梯度截断、P3 调制方式不匹配）的解决方案详见 v1（SimVQ性能分析报告.docx）第四~五章（方案 B）。补充问题 N1（验证管线使用错误信道）详见 v3（SimVQ_遗漏问题补充报告_v3.docx）新问题 1。"
            )
            for run in ref.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 100, 180)
                run.italic = True
            break

    for i, para in enumerate(doc.paragraphs):
        if "第三部分：SimVQ 模块深度解析" in para.text:
            ref = para.insert_paragraph_before(
                "▸ 相关内容：本部分分析的 BN 失效问题在 v1 中未涉及，是本报告的独有贡献。补充问题 N3（Adam β₁=0.5 加剧训练不稳定性）与本部分的优化器行为分析相关，详见 v3 新问题 3。"
            )
            for run in ref.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 100, 180)
                run.italic = True
            break

    for i, para in enumerate(doc.paragraphs):
        if "第四部分：逐模块微观问题清单" in para.text:
            ref = para.insert_paragraph_before(
                "▸ 相关内容：本部分分析的 SimVQ 局限性在 v1 中简要提及（P5 码本容量不足），本报告给出了严格的数学论证。补充问题 N4（编码器/码本尺度不匹配）、N5（commitment_cost 不对称）、N8（投影层无正则化）与本部分直接相关，详见 v3。"
            )
            for run in ref.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 100, 180)
                run.italic = True
            break

    for i, para in enumerate(doc.paragraphs):
        if "第五部分：快速验证方法论" in para.text:
            ref = para.insert_paragraph_before(
                "▸ 相关内容：上一部分的微观问题已在 v1 的方案和 v3 的补充问题中分散覆盖。本部分的损失函数和数据处理分析在 v1 的 P4（MSE损失）和 P6（数据域偏移）中有对应的高层概述。"
            )
            for run in ref.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 100, 180)
                run.italic = True
            break

    # 本部分在末尾处添加一下总结性的交叉引用
    for i, para in enumerate(doc.paragraphs):
        if "附录：关键问题与代码位置速查表" in para.text:
            ref = para.insert_paragraph_before(
                "▸ 相关内容：快速验证方法论的各项代理指标（指标1-6）与 v1 方案 B 的各项改动直接对应——指标4（鲁棒性差距）衡量改动1（可微分AWGN信道）的效果，指标5（梯度比）衡量改动4（STE修复）的效果，指标2（码本利用率）衡量改动6（扩大码本）的效果。"
            )
            for run in ref.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 100, 180)
                run.italic = True
            break

    # 3. 在末尾添加"相关报告"章节
    add_related_reports_section(doc, "v2")

    doc.save(path)
    print("v2 处理完成 ✓")


def process_v3():
    """处理 v3 报告"""
    path = f"{BASE}/SimVQ_遗漏问题补充报告_v3.docx"
    print(f"处理 {path}...")
    doc = Document(path)

    # 1. 更新"关于本报告"段落（在原有基础上增强）
    for i, para in enumerate(doc.paragraphs):
        if "本报告是前两份报告" in para.text:
            # 在其后追加补充说明
            note = add_paragraph_after(
                doc,
                "本报告是前两份报告",
                "三份报告的定位与关系：v1（SimVQ性能分析报告.docx）提供项目概览、8个核心问题诊断和3套优化方案；v2（SimVQ_深度性能分析报告_v2.docx）提供深度技术剖析（BN/GN分析、SimVQ数学局限性、快速验证方法论）；本报告（v3）补充前两份未覆盖的11个新问题，并提供33问题总索引。建议按 v1->v2->v3 顺序阅读。",
                italic=True
            )
            if note:
                for run in note.runs:
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(100, 100, 100)
            break

    # 2. 在每个新问题末尾添加"关联问题"标注
    cross_refs = {
        "N1": "关联问题：v1 P1（训练/测试信道不一致）— N1 是该问题的延伸，揭示了验证阶段同样受错误信道影响，导致最优模型从未被正确选出。另见 v2 第一部分（双信道模型逐层对比）。",
        "N2": "关联问题：v1 P4（仅用 MSE 损失）— N2 与 P4 共同指向重建质量问题，N2 额外揭示了当前 MS-SSIM/PSNR 报告值系统性偏高。另见 v2 4.4 节（损失函数设计问题）和问题 L1。",
        "N3": "关联问题：v2 2.2 节（BN 失效模式3：batch_size 敏感性）— β₁=0.5 放大了 batch_size=24 时的梯度噪声。另见 v2 4.3 节（信道模块设计问题）中关于训练不稳定性的讨论。",
        "N4": "关联问题：v1 P5（码本容量不足）+ v2 第三部分（SimVQ 深度解析）— N4 从尺度匹配角度补充了 P5 和 v2 的码本分析，提供了 L2 归一化的具体解决方案。另见 v2 3.2 节（冻结嵌入+投影层的数学局限性）。",
        "N5": "关联问题：v1 P7（VQ 损失权重失衡）— N5 和 P7 共同指向 VQ 训练的优化不平衡，N5 从 commitment_cost 角度补充了 P7 的 LAYER_LOSS_WEIGHTS 分析。另见 v2 4.4 节（损失函数设计问题）和问题 L2。",
        "N6": "关联问题：v2 1.2 节（FiniteBlocklengthChannel 缺陷5：Clamping 码字偏差）— 两者都涉及索引恢复偏差，N6 是比特层面的对应问题。另见 v2 问题 P12。",
        "N7": "关联问题：v2 1.3 节（STE 梯度流分析）+ 指标5（梯度 SNR）— N7 的发现印证了 v2 中关于编解码器梯度不平衡的分析。",
        "N8": "关联问题：v2 3.4 节（码本坍缩风险）+ 3.2 节（投影层有效自由度分析）— N8 提供了具体的正则化方案来应对 v2 中识别的投影坍缩风险。",
        "N9": "关联问题：v1 P6（训练/测试数据域偏移）— N9 从预处理细节角度补充了 P6，揭示了宽高比失真这一具体问题。另见 v2 4.5 节（数据处理问题）和问题 D1、D2。",
        "N10": "关联问题：v2 第五部分（快速验证方法论）- N10 关注的是恢复训练后 LR 过低导致无法判断改动有效性的实操问题，与 v2 Phase 2 中建议的[使用2-3x正常学习率]相关。",
        "N11": "关联问题：v1 P4（MSE vs MS-SSIM 指标不一致）— N11 从评估函数鲁棒性角度补充了 P4，指出少数边界情况图像可能不合理地拉低平均 MS-SSIM。",
    }

    for problem_id, ref_text in cross_refs.items():
        # Find paragraph containing the problem header
        search_terms = {
            "N1": "新问题 1：",
            "N2": "新问题 2：",
            "N3": "新问题 3：",
            "N4": "新问题 4：",
            "N5": "新问题 5：",
            "N6": "新问题 6：",
            "N7": "新问题 7：",
            "N8": "新问题 8：",
            "N9": "新问题 9：",
            "N10": "新问题 10：",
            "N11": "新问题 11：",
        }

        search = search_terms.get(problem_id)
        if not search:
            continue

        # Find the end of this problem's section (where the next problem or section starts)
        next_markers = [f"新问题 {j}：" for j in range(1, 12) if j > int(problem_id[1:])]
        next_markers.append("新问题汇总")
        next_markers.append("三版报告问题总索引")

        problem_start_idx = None
        problem_end_idx = None

        for i, para in enumerate(doc.paragraphs):
            if search in para.text and problem_start_idx is None:
                problem_start_idx = i
            elif problem_start_idx is not None:
                for marker in next_markers:
                    if marker in para.text:
                        problem_end_idx = i
                        break
                if problem_end_idx:
                    break

        if problem_start_idx:
            # Insert the cross-ref before the next problem or at the end of this problem
            insert_idx = problem_end_idx if problem_end_idx else problem_start_idx + 5
            if insert_idx < len(doc.paragraphs):
                target_para = doc.paragraphs[insert_idx]
                ref_para = target_para.insert_paragraph_before("")
                ref_para2 = target_para.insert_paragraph_before(
                    f"▸ {ref_text}"
                )
                for run in ref_para2.runs:
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0, 100, 180)
                    run.italic = True

    # 3. 在"三版报告问题总索引"后添加"推荐阅读顺序"
    for i, para in enumerate(doc.paragraphs):
        if "三版报告问题总索引" in para.text:
            # Find end of the index table, then add reading guide
            # We'll add it after the last ID entry
            pass

    # Find the end of the document for reading guide
    found_last = False
    for i, para in enumerate(doc.paragraphs):
        if "N11" in para.text and "MS-SSIM失败返回0.0" in para.text:
            found_last = True
            continue
        if found_last and para.text.strip() == "":
            # Skip blank lines after the last entry
            continue
        if found_last and para.text.strip() and "N11" not in para.text:
            # This is after the last N11 entry, insert reading guide here
            doc.paragraphs[i].insert_paragraph_before("")
            guide_title = doc.paragraphs[i].insert_paragraph_before("推荐阅读顺序")
            if guide_title.runs:
                guide_title.runs[0].bold = True
                guide_title.runs[0].font.size = Pt(14)

            guide_items = [
                "1. 首先阅读 v1（SimVQ性能分析报告.docx）— 了解项目背景、系统架构和 8 个核心问题，建立全局认知。重点关注：CRITICAL 级别的 P1-P3（信道不一致、STE梯度、调制不匹配）以及方案 B 的 5 项改动。",
                "2. 然后阅读 v2（SimVQ_深度性能分析报告_v2.docx）— 深入理解 BN 的 6 个失效模式（第二部分）、SimVQ 的数学局限性（第三部分）、以及快速验证方法论（第五部分）。重点关注：GN 替换方案（2.3节）、投影坍缩监控（3.4节）、Phase 2 快速验证流程。",
                "3. 最后阅读本报告（v3）— 了解前两份报告遗漏的 11 个问题及其修正方向。重点关注：N1（验证管线修复，CRITICAL）、N2（Tanh 输出约束，HIGH）、N3（Adam β₁ 恢复，HIGH）、N4（L2 归一化，HIGH）。",
                "4. 实施时，以 v1 方案 B 为基础框架，融入 v3 的 N1-N4 修正，采用 v2 第五部分的快速验证方法论分阶段验证，使用 v3 的 33 问题总索引作为检查清单逐项核对。",
            ]
            for item in guide_items:
                p = doc.paragraphs[i].insert_paragraph_before(item)
                for run in p.runs:
                    run.font.size = Pt(10)
            break

    # 4. 添加"相关报告"章节
    add_related_reports_section(doc, "v3")

    doc.save(path)
    print("v3 处理完成 ✓")


if __name__ == "__main__":
    process_v1()
    process_v2()
    process_v3()
    print("\n全部三个文件处理完成！")
